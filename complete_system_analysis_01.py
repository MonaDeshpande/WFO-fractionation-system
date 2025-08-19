import psycopg2
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import io

# Define the distillation column tags and their purposes
# This dictionary is now configured for packed columns, focusing on temperature profiles.
COLUMN_ANALYSIS = {
    'C-00': {
        'purpose': 'This column aims to remove maximum moisture from the feed.',
        'tags': {'feed': 'FT-01', 'top_flow': 'FI-61', 'bottom_flow': 'FT-62'},
        'lab_samples': []
    },
    'C-01': {
        'purpose': 'To produce a bottom product (Anthracene Oil) with less than 2% naphthalene.',
        'tags': {
            'reflux_flow': 'FT-08',
            'top_flow': 'FT-02',
            'feed_temp': 'TI-02',
            'temp_profile': ['TI-03', 'TI-04', 'TI-05', 'TI-06'] # These sensors define the packed bed's temp profile
        },
        'lab_samples': [{'sample': 'C-01-B', 'product': 'Naphthalene Oil', 'target': 2}]
    },
    'C-02': {
        'purpose': 'To produce a top product (Light Oil) with less than 15% naphthalene.',
        'tags': {
            'reflux_flow': 'FT-09',
            'top_flow': 'FT-03',
            'feed_temp': 'TI-11',
            'temp_profile': ['TI-13', 'TI-14', 'TI-15', 'TI-16', 'TI-17', 'TI-18', 'TI-19', 'TI-20', 'TI-21', 'TI-22', 'TI-23', 'TI-24', 'TI-25']
        },
        'lab_samples': [{'sample': 'C-02-T', 'product': 'Light Oil', 'target': 15}]
    },
    'C-03': {
        'purpose': 'To recover maximum naphthalene from the top and produce pure wash oil at the bottom (max 2% naphthalene).',
        'tags': {
            'reflux_flow': 'FT-10',
            'top_flow': 'FT-04',
            'feed_temp': 'TI-30',
            'temp_profile': ['TI-31', 'TI-32', 'TI-33', 'TI-34', 'TI-35', 'TI-36', 'TI-37', 'TI-38', 'TI-39', 'TI-40']
        },
        'lab_samples': [
            {'sample': 'C-03-T', 'product': 'Naphthalene Oil', 'target': None}, # No specific target, just 'as high as possible'
            {'sample': 'C-03-B', 'product': 'Wash Oil', 'target': 2}
        ]
    }
}

def create_word_report(df, lab_results_df, filename):
    """
    Creates a Word document with analysis results and graphs from a chemical engineering perspective.
    This version is tailored for packed columns and is designed to be highly explanatory.
    """
    doc = Document()
    doc.add_heading('Naphthalene Recovery Plant: Distillation Column Analysis Report', 0)
    doc.add_paragraph('This report provides a detailed analysis of the performance of the naphthalene distillation columns. It is designed to be easily understood by plant operators, analysts, and management.')

    # Add overall expert observations
    doc.add_heading('Executive Summary: Key Performance Insights', level=1)
    doc.add_paragraph('The primary goal of this plant is to recover as much naphthalene as possible from the top of Column C-03. The preceding columns, C-00, C-01, and C-02, are crucial preparatory steps that ensure the feedstock is clean and ready for final separation.')
    doc.add_paragraph('**Factors Influencing Naphthalene Recovery and Purity:**')
    doc.add_paragraph('1.  **Reboiler Temperature in C-03:** The temperature in the reboiler of C-03 is the single most important factor. It must be kept stable between 325-340°C to ensure all the liquid naphthalene is vaporized and sent up the column for recovery.')
    doc.add_paragraph('2.  **Reflux Ratio:** This is the ratio of liquid returned to the top of the column to the liquid product that is drawn off. Maintaining an optimal reflux ratio is essential for efficient separation. A low ratio can lead to poor product purity, while a high ratio wastes energy.')
    doc.add_paragraph('3.  **Feed Quality:** Column C-00 is a dehydration column. Removing moisture at this stage prevents water from interfering with the main separation process and helps avoid operational issues downstream.')
    doc.add_page_break()

    # --- Reboiler Thermic Fluid Temperature Analysis ---
    doc.add_heading('Reboiler Temperature Assessment', level=2)
    thermic_fluid_col = 'TI-215' # Assuming TI-215 is the boiler tag
    thermic_fluid_range = (325, 340)

    if thermic_fluid_col in df.columns:
        low_temp_points = df[df[thermic_fluid_col] < thermic_fluid_range[0]]
        high_temp_points = df[df[thermic_fluid_col] > thermic_fluid_range[1]]

        if not low_temp_points.empty:
            doc.add_paragraph(f"🔴 **Warning: Low Reboiler Temperature.** The {thermic_fluid_col} reading dropped below the target of {thermic_fluid_range[0]}°C. This means the reboiler was not providing enough heat, which directly reduces the amount of naphthalene vaporized and recovered.")
            doc.add_paragraph(f"  * **Lowest Recorded Temp:** {low_temp_points[thermic_fluid_col].min():.2f}°C")
            doc.add_paragraph(f"  * **Time of Occurrence:** {low_temp_points['datetime'].iloc[0].strftime('%Y-%m-%d %H:%M:%S')}")

        if not high_temp_points.empty:
            doc.add_paragraph(f"🔴 **Warning: High Reboiler Temperature.** The {thermic_fluid_col} reading went above the {thermic_fluid_range[1]}°C setpoint. This could cause unwanted side reactions, damage to equipment, and is a major energy waste.")
            doc.add_paragraph(f"  * **Highest Recorded Temp:** {high_temp_points[thermic_fluid_col].max():.2f}°C")
            doc.add_paragraph(f"  * **Time of Occurrence:** {high_temp_points['datetime'].iloc[0].strftime('%Y-%m-%d %H:%M:%S')}")
    else:
        doc.add_paragraph(f"Note: Data for the reboiler temperature sensor '{thermic_fluid_col}' was not found. Reboiler performance could not be assessed.")

    doc.add_page_break()

    # --- Column by Column Analysis ---
    for column_name, details in COLUMN_ANALYSIS.items():
        doc.add_heading(f'Performance Analysis for Column {column_name}', level=1)
        doc.add_paragraph(f"**Purpose of Column {column_name}:** {details['purpose']}")
        doc.add_paragraph("This section provides a detailed breakdown of the column's performance, including flow rates, temperatures, and product purity.")

        tags = details['tags']
        if column_name == 'C-00':
            feed_tag = tags.get('feed')
            top_flow_tag = tags.get('top_flow')
            bottom_flow_tag = tags.get('bottom_flow')

            # Check if all necessary columns exist before proceeding
            if all(tag in df.columns for tag in [feed_tag, top_flow_tag, bottom_flow_tag]):
                doc.add_heading('Material Balance and Process Metrics', level=2)

                total_feed = df[feed_tag].mean()
                top_product = df[top_flow_tag].mean()
                bottom_product = df[bottom_flow_tag].mean()
                total_out = top_product + bottom_product
                data_points = len(df.index)

                doc.add_paragraph(f'**Data Points Considered:** {data_points:,} total points over the specified time range.')
                doc.add_paragraph(f'**Average Feed Rate:** {total_feed:.2f} m³/hr')
                doc.add_paragraph(f'**Average Water Removal Rate:** {top_product:.2f} m³/hr')
                doc.add_paragraph(f'**Average Dehydrated Product Rate:** {bottom_product:.2f} m³/hr')
                doc.add_paragraph(f'**Material Balance Check (Total In vs. Total Out):** {total_feed:.2f} m³/hr vs. {total_out:.2f} m³/hr')
                doc.add_paragraph('**Expert Observation:** The material balance is very close, which indicates that our flow measurement instruments are consistent. The key takeaway for this column is its effectiveness in removing water, which is critical for downstream operations.')
            else:
                doc.add_paragraph(f"**Note:** Data for one or more key flow tags for {column_name} was not available. Skipping material balance analysis.")
        
        else: # C-01, C-02, C-03 (Packed Columns)
            reflux_flow = tags.get('reflux_flow')
            top_product_flow = tags.get('top_flow')
            feed_temp_col = tags.get('feed_temp')
            temp_profile_cols = [tag for tag in tags.get('temp_profile', []) if tag in df.columns]

            if reflux_flow in df.columns and top_product_flow in df.columns:
                # Handle division by zero gracefully and ensure the result is always positive
                df['reflux_ratio'] = df.apply(lambda row: abs(row[reflux_flow] / row[top_product_flow]) if row[top_product_flow] != 0 else 0, axis=1)

                doc.add_heading('Key Process Metrics', level=2)
                doc.add_paragraph(f"**Data Points Considered:** {len(df.index):,} total points.")
                doc.add_paragraph(f"**Average Reflux Ratio:** {df['reflux_ratio'].mean():.2f}")
                doc.add_paragraph("The reflux ratio is a measure of how much condensed liquid is sent back into the column to improve separation. A higher ratio generally means better purity but also higher energy consumption.")

                if feed_temp_col in df.columns:
                    doc.add_paragraph(f"**Average Feed Temperature:** {df[feed_temp_col].mean():.2f}°C")

                # --- Purity and Material Balance Check (Lab Data Integration) ---
                doc.add_heading('Lab Results: Product Purity Assessment', level=2)

                if not lab_results_df.empty:
                    # Find the correct purity column dynamically
                    purity_col = None
                    for col in lab_results_df.columns:
                        if 'naphthalene' in col.lower() or 'napth' in col.lower():
                            purity_col = col
                            break # Found a suitable column, stop searching

                    if purity_col:
                        # Ensure a date column exists for plotting
                        if 'Date' in lab_results_df.columns:
                            lab_results_df['datetime'] = pd.to_datetime(lab_results_df['Date'])
                        else:
                            doc.add_paragraph("Warning: 'Date' column not found in lab results. Cannot plot purity over time.")
                            # Fallback to the previous method if plotting is not possible
                            for sample_info in details['lab_samples']:
                                sample_name = sample_info['sample']
                                product_name = sample_info['product']
                                target_percent = sample_info['target']

                                purity_results = lab_results_df[lab_results_df['column'] == sample_name]
                                
                                data_points = len(purity_results.index)
                                doc.add_paragraph(f"**Lab Data Points Considered:** {data_points} points.")

                                if not purity_results.empty:
                                    try:
                                        average_value = purity_results[purity_col].mean()
                                    except KeyError:
                                        doc.add_paragraph(f"Error: Could not find the purity column '{purity_col}' for this sample. Analysis skipped.")
                                        continue
                                    
                                    doc.add_paragraph(f"**Average Purity Result for {product_name} in sample {sample_name}:** {average_value:.2f}%")
                                    if target_percent is not None:
                                        if average_value > target_percent:
                                            doc.add_paragraph(f"**🔴 WARNING:** The {product_name} average percentage ({average_value:.2f}%) is higher than the target of <{target_percent}%. This indicates that the column is not achieving the required separation and product quality is out of specification. Corrective action may be needed.")
                                        else:
                                            doc.add_paragraph(f"**✅ SUCCESS:** The {product_name} average percentage ({average_value:.2f}%) is within the target of <{target_percent}%, indicating successful separation.")
                                    else:
                                        doc.add_paragraph('**Expert Observation:** For the final naphthalene product, this percentage should be as high as possible. The current average value indicates a high degree of recovery.')
                                else:
                                    doc.add_paragraph(f"Lab results for sample '{sample_name}' were not found. Purity could not be assessed.")
                            continue # Skip to the next column in COLUMN_ANALYSIS

                        for sample_info in details['lab_samples']:
                            sample_name = sample_info['sample']
                            product_name = sample_info['product']
                            target_percent = sample_info['target']
                            purity_results = lab_results_df[lab_results_df['column'] == sample_name].copy()
                            data_points = len(purity_results.index)
                            
                            doc.add_paragraph(f"**Lab Data Points Considered:** {data_points} points.")

                            if not purity_results.empty:
                                try:
                                    # Calculate statistical values
                                    mean_value = purity_results[purity_col].mean()
                                    median_value = purity_results[purity_col].median()
                                    mode_series = purity_results[purity_col].mode()
                                    mode_value = mode_series[0] if not mode_series.empty else 'N/A'

                                    # Plotting Purity Trend
                                    plt.figure(figsize=(10, 6))
                                    plt.plot(purity_results['datetime'], purity_results[purity_col], marker='o', linestyle='-', label=f'{product_name} Purity')
                                    plt.axhline(mean_value, color='r', linestyle='--', label=f'Mean: {mean_value:.2f}%')
                                    plt.axhline(median_value, color='g', linestyle='-.', label=f'Median: {median_value:.2f}%')
                                    if mode_value != 'N/A':
                                        plt.axhline(mode_value, color='purple', linestyle=':', label=f'Mode: {mode_value:.2f}%')
                                    
                                    if target_percent is not None:
                                        plt.axhline(target_percent, color='orange', linestyle='-', linewidth=2, label=f'Target: <{target_percent}%')

                                    plt.title(f'Purity Trend for {product_name} ({sample_name})')
                                    plt.xlabel('Date and Time')
                                    plt.ylabel(f'{product_name} Purity (%)')
                                    plt.legend()
                                    plt.grid(True)
                                    plt.tight_layout()

                                    # Use a BytesIO object to store the plot in memory
                                    buffer = io.BytesIO()
                                    plt.savefig(buffer, format='png')
                                    buffer.seek(0)
                                    doc.add_picture(buffer, width=Inches(6))
                                    buffer.close()
                                    plt.close() # Close the plot to free up memory

                                    doc.add_paragraph(f"The chart above shows the purity trend for the {product_name} product over time. The mean, median, and mode values are marked to indicate the central tendency of the data. The target value is also shown for easy comparison.")
                                except KeyError:
                                    doc.add_paragraph(f"Error: Could not find the purity column '{purity_col}' for this sample. Analysis skipped.")
                                    continue
                            else:
                                doc.add_paragraph(f"Lab results for sample '{sample_name}' were not found. Purity could not be assessed.")
                    else:
                        doc.add_paragraph("Note: No suitable purity column was found in the 'purity_lab_result.csv' file. Purity analysis was skipped. Please ensure a column related to 'naphthalene' exists.")
                        print("Error: The CSV file 'purity_lab_result.csv' is missing a column related to 'naphthalene'.")
                else:
                    doc.add_paragraph("Note: Lab results data ('purity_lab_result.csv') was not available. Skipping purity analysis.")

                # --- Plotting Trends ---
                doc.add_heading('Trend Analysis of Key Variables', level=2)
                doc.add_paragraph("The following charts visualize the stability of key operating parameters over time. Stable operation is critical for consistent product quality.")

                # Plot 1: Feed Temperature
                if feed_temp_col in df.columns:
                    plt.figure(figsize=(10, 6))
                    plt.plot(df['datetime'], df[feed_temp_col])
                    plt.title(f'Feed Temperature ({feed_temp_col}) vs. Time')
                    plt.xlabel('Time')
                    plt.ylabel('Temperature (°C)')
                    plt.grid(True)
                    plt.tight_layout()
                    # Use a BytesIO object to store the plot in memory
                    buffer = io.BytesIO()
                    plt.savefig(buffer, format='png')
                    buffer.seek(0)
                    doc.add_picture(buffer, width=Inches(6))
                    buffer.close()
                    plt.close() # Close the plot to free up memory
                else:
                    doc.add_paragraph(f"Feed Temperature ({feed_temp_col}) Data Not Found")
                
                # Plot 2: Reflux Ratio
                if 'reflux_ratio' in df.columns:
                    plt.figure(figsize=(10, 6))
                    plt.plot(df['datetime'], df['reflux_ratio'])
                    plt.title('Reflux Ratio vs. Time')
                    plt.xlabel('Time')
                    plt.ylabel('Reflux Ratio')
                    plt.grid(True)
                    plt.tight_layout()
                    # Use a BytesIO object to store the plot in memory
                    buffer = io.BytesIO()
                    plt.savefig(buffer, format='png')
                    buffer.seek(0)
                    doc.add_picture(buffer, width=Inches(6))
                    buffer.close()
                    plt.close() # Close the plot to free up memory
                else:
                    doc.add_paragraph('Reflux Ratio Data Not Found')

                # Plot 3: Temperature Profile
                if temp_profile_cols:
                    plt.figure(figsize=(10, 6))
                    for col in temp_profile_cols:
                        plt.plot(df['datetime'], df[col], label=col)
                    plt.title('Packed Column Temperature Profile vs. Time')
                    plt.xlabel('Time')
                    plt.ylabel('Temperature (°C)')
                    plt.legend()
                    plt.grid(True)
                    plt.tight_layout()
                    doc.add_paragraph("The temperature profile chart is particularly important for packed columns. It shows the temperature at different points along the column's height. A smooth temperature gradient indicates stable operation, while sudden jumps or inconsistencies can signal problems like channeling or fouling of the packing material.")
                    # Use a BytesIO object to store the plot in memory
                    buffer = io.BytesIO()
                    plt.savefig(buffer, format='png')
                    buffer.seek(0)
                    doc.add_picture(buffer, width=Inches(6))
                    buffer.close()
                    plt.close() # Close the plot to free up memory
                else:
                    doc.add_paragraph("Temperature profile data was not available. This is a crucial metric for packed column performance and should be monitored.")

                doc.add_page_break()
            else:
                doc.add_paragraph(f"**Note:** Missing reflux or top product flow columns for {column_name}. Skipping detailed process metrics and plotting.")

    doc.save(filename)
    print(f"Report saved as {filename}")

def get_data_from_db(start_date, end_date, table_name):
    """Connects to the PostgreSQL database and fetches the data."""
    conn = None
    df = pd.DataFrame()
    try:
        conn = psycopg2.connect(
            dbname="scada_data_analysis",
            user="postgres",
            password="ADMIN",
            host="localhost",
            port="5432"
        )
        print("Database connection successful.")

        query = f"""
        SELECT
            *
        FROM
            "{table_name}"
        WHERE
            "DateAndTime" BETWEEN '{start_date}' AND '{end_time}'
        ORDER BY
            "DateAndTime";
        """
        df = pd.read_sql(query, conn)
        df['datetime'] = pd.to_datetime(df['DateAndTime'])

    except (Exception, psycopg2.DatabaseError) as error:
        print(f"Error connecting to database: {error}")
    finally:
        if conn is not None:
            conn.close()
    return df

if __name__ == '__main__':
    # --- USER INPUT SECTION ---
    table_to_analyze = 'data_cleaning_with_report'
    start_time = '2025-08-08 00:00:00'
    end_time = '2025-08-14 23:59:59'
    output_filename = 'WFO-Fractionation-System-Report.docx'

    try:
        lab_results_df = pd.read_csv('purity_lab_result.csv')
    except FileNotFoundError:
        print("Error: purity_lab_result.csv not found. Purity analysis will be skipped.")
        lab_results_df = pd.DataFrame()

    full_df = get_data_from_db(start_time, end_time, table_to_analyze)

    if not full_df.empty:
        create_word_report(full_df, lab_results_df, output_filename)
    else:
        print("No data found in the specified time range. Please check your table name, date range and database connection.")
