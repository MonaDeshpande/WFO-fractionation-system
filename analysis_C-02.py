import pandas as pd
from sqlalchemy import create_engine
import psycopg2
from datetime import datetime
import matplotlib.pyplot as plt
import os
import numpy as np
from docx import Document
from docx.shared import Inches
import sqlalchemy
import re

# Database connection parameters (update with your actual details)
DB_HOST = "localhost"
DB_NAME = "scada_data_analysis"
DB_USER = "postgres"
DB_PASSWORD = "ADMIN"

# Define units for each tag
TAG_UNITS = {
    'FT-01': 'kg/h',
    'FT-02': 'kg/h',
    'FT-03': 'kg/h',
    'FT-06': 'kg/h',
    'FT-61': 'kg/h',
    'FI-103': 'kg/h',
    'FI-202': 'kg/h',
    'TI-11': 'degC',
    'TI-13': 'degC',
    'TI-14': 'degC',
    'TI-15': 'degC',
    'TI-16': 'degC',
    'TI-17': 'degC',
    'TI-18': 'degC',
    'TI-19': 'degC',
    'TI-20': 'degC',
    'TI-21': 'degC',
    'TI-22': 'degC',
    'TI-23': 'degC',
    'TI-24': 'degC',
    'TI-25': 'degC',
    'TI-26': 'degC',
    'TI-27': 'degC',
    'TI-28': 'degC',
    'TI-29': 'degC',
    'TI-30': 'degC',
    'TI-72A': 'degC',
    'TI-72B': 'degC',
    'PTT-02': 'mmHg',
    'PTB-02': 'mmHg',
    'DIFFERENTIAL_PRESSURE': 'mmHg',
    'LI-03': '%',
    'REBOILER_HEAT_DUTY': 'kW',
    'CONDENSER_HEAT_DUTY': 'kW',
    'REFLUX_RATIO': '',
    'MATERIAL_BALANCE_ERROR': '%',
    'NAPHTHALENE_LOSS_PERCENTAGE': '%'
}

# File paths for saving generated plots and report
output_report_path = "C-02_Analysis_Report.docx"
output_temp_plot_path = "C-02_Temperature_Profile.png"
output_dp_plot_path = "C-02_Differential_Pressure.png"
output_trends_plot_path = "C-02_Daily_Trends.png"

# Engineering constants for heat duty calculations
THERMIC_FLUID_SPECIFIC_HEAT = 2.0  # kJ/(kg·°C) - Assumed value
WATER_SPECIFIC_HEAT = 4.186       # kJ/(kg·°C)

def connect_to_database():
    """Establishes a connection to the PostgreSQL database."""
    try:
        engine = create_engine(f'postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}/{DB_NAME}')
        print("Database connection successful.")
        return engine
    except Exception as e:
        print(f"Error connecting to the database: {e}")
        return None

def get_scada_data(engine):
    """Retrieves specific SCADA data for the C-02 column and related streams."""
    try:
        desired_columns = [
            "DateAndTime", "FT-01", "FT-02", "FT-03", "FT-06", "FT-09", "FT-61", "TI-11", "TI-13", "TI-14", "TI-15", "TI-16", 
            "TI-17", "TI-18", "TI-19", "TI-20", "TI-21", "TI-22", "TI-23", "TI-24", "TI-25", "TI-26",
            "TI-27", "TI-28", "TI-29", "TI-30", "TI-72A", "TI-72B", "PTT-02", "PTB-02", "LI-03",
            "FI-103", "FI-202"
        ]
        
        inspector = sqlalchemy.inspect(engine)
        columns = inspector.get_columns('wide_scada_data')
        column_names = [col['name'] for col in columns]
        
        final_columns = []
        for d_col in desired_columns:
            for db_col in column_names:
                if d_col.lower().replace('-', '') == db_col.lower().replace('-', ''):
                    final_columns.append(f'"{db_col}"')
                    break
        
        if not final_columns:
            print("Error: No matching columns found. Data retrieval failed.")
            return None

        select_clause = ", ".join(final_columns)
        query = f"""
        SELECT {select_clause}
        FROM wide_scada_data
        WHERE "DateAndTime" BETWEEN '2025-08-08 00:00:00' AND '2025-08-20 23:59:59'
        ORDER BY "DateAndTime";
        """
        
        df = pd.read_sql(query, engine)
        df['DateAndTime'] = pd.to_datetime(df['DateAndTime'])
        print("SCADA data for C-02 and related streams retrieved successfully.")
        return df
    except Exception as e:
        print(f"Error retrieving SCADA data: {e}")
        return None

def get_composition_data():
    """
    Simulates reading composition data from a lab analysis report.
    Returns a dictionary of compositions for each component at specific sample points.
    """
    try:
        # Simulate data from the provided description
        composition_data = {
            'Naphthalene': {
                'P-01': 0.1,    # Assumed composition in total feed (10%)
                'C-01-B': 0.02, # Naphthalene loss in C-01 bottom product (2% is acceptable)
                'C-02-T': 0.08  # Assumed composition in C-02 top product (8%)
            },
            'Thianaphthene': {
                'P-01': 0.05,
                'C-01-B': 0.01,
                'C-02-T': 0.04
            },
            'Quinoline': {
                'P-01': 0.03,
                'C-01-B': 0.02,
                'C-02-T': 0.01
            },
            'Unknown Impurity': {
                'P-01': 0.01,
                'C-01-B': 0.005,
                'C-02-T': 0.002
            },
            'Moisture': {
                'P-01': 0.15 # 15% moisture in the total feed
            }
        }
        return composition_data
    except Exception as e:
        print(f"Error simulating composition data: {e}. Using default values.")
        return None

def perform_analysis(df):
    """
    Performs key calculations for the process, including material/energy balances
    and component-wise analysis.
    """
    if df is None or df.empty:
        return {}, df, {}

    analysis_results = {}
    
    # Material Balance Analysis for the entire process
    if all(tag in df.columns for tag in ['FT-01', 'FT-03', 'FT-06', 'FT-61']):
        total_feed_avg = df['FT-01'].mean()
        moisture_removed_avg = df['FT-61'].mean()
        top_product_flow_avg = df['FT-03'].mean()
        bottom_product_flow_avg = df['FT-06'].mean()
        
        net_feed = total_feed_avg - moisture_removed_avg
        total_products = top_product_flow_avg + bottom_product_flow_avg
        
        if net_feed > 0:
            overall_balance_error = ((net_feed - total_products) / net_feed) * 100
            analysis_results['Overall Material Balance Error (%)'] = abs(overall_balance_error)

    # Component-wise Material Balance
    composition_data = get_composition_data()
    if composition_data:
        analysis_results['Component-wise Material Balance'] = {}
        
        # Assume FT-02 is the feed to C-02, which is the top product of C-01
        c02_feed_flow_avg = df['FT-02'].mean()
        
        for component, comps in composition_data.items():
            if component == 'Moisture':
                # Moisture is removed in the first step
                input_mass_flow = df['FT-01'].mean() * comps['P-01']
                output_mass_flow = df['FT-61'].mean()
            else:
                input_mass_flow = df['FT-01'].mean() * comps['P-01']
                # The user description is a bit ambiguous, but we will assume
                # the component leaves in either the C-02 top or C-01 bottom.
                output_mass_flow = (df['FT-03'].mean() * comps['C-02-T']) + (df['FT-06'].mean() * comps['C-01-B'])
            
            if input_mass_flow > 0:
                component_balance_error = ((input_mass_flow - output_mass_flow) / input_mass_flow) * 100
                analysis_results['Component-wise Material Balance'][f'Error for {component}'] = abs(component_balance_error)
                
    # Naphthalene Loss & Impurity Analysis
    if composition_data and 'Naphthalene' in composition_data:
        top_prod_comp = composition_data['Naphthalene'].get('C-02-T')
        c01_bottom_comp = composition_data['Naphthalene'].get('C-01-B')
        
        if top_prod_comp is not None:
            analysis_results['Naphthalene in Top Product (%)'] = top_prod_comp * 100
            if top_prod_comp > 0.15: # 15% threshold for C-02
                analysis_results['C-02 Naphthalene Loss Status'] = "ALERT: Naphthalene loss is above 15%."
            else:
                analysis_results['C-02 Naphthalene Loss Status'] = "Naphthalene loss is within acceptable limits (below 15%)."

        if c01_bottom_comp is not None:
            analysis_results['Naphthalene in C-01 Bottom Product (%)'] = c01_bottom_comp * 100
            if c01_bottom_comp > 0.02: # 2% threshold for C-01 bottom
                analysis_results['C-01 Naphthalene Loss Status'] = "ALERT: Naphthalene loss is above 2%."
            else:
                analysis_results['C-01 Naphthalene Loss Status'] = "Naphthalene loss is within acceptable limits (below 2%)."

    # Reflux Ratio
    if 'FT-09' in df.columns and 'FT-03' in df.columns:
        reflux_flow_avg = df['FT-09'].mean()
        top_product_flow_avg = df['FT-03'].mean()
        
        if top_product_flow_avg > 0:
            reflux_ratio = reflux_flow_avg / top_product_flow_avg
            analysis_results['Average Reflux Ratio'] = reflux_ratio
        else:
            analysis_results['Average Reflux Ratio'] = "N/A (Zero Top Product Flow)"
            
    # Differential Pressure (DP) Calculation
    if 'PTT-02' in df.columns and 'PTB-02' in df.columns:
        df['DIFFERENTIAL_PRESSURE'] = df['PTB-02'] - df['PTT-02']
        analysis_results['Average Differential Pressure'] = df['DIFFERENTIAL_PRESSURE'].mean()
        analysis_results['Maximum Differential Pressure'] = df['DIFFERENTIAL_PRESSURE'].max()
        
    # Energy Balance
    # Reboiler Heat Duty
    if all(tag in df.columns for tag in ['TI-72A', 'TI-72B', 'FI-202']):
        df['REBOILER_HEAT_DUTY'] = df['FI-202'] * THERMIC_FLUID_SPECIFIC_HEAT * (df['TI-72B'] - df['TI-72A'])
        analysis_results['Average Reboiler Heat Duty'] = df['REBOILER_HEAT_DUTY'].mean()

    # Condenser Heat Duty (Main)
    if all(tag in df.columns for tag in ['TI-26', 'FI-103']):
        df['CONDENSER_HEAT_DUTY'] = df['FI-103'] * WATER_SPECIFIC_HEAT * (df['TI-26'] - df['TI-11'])
        analysis_results['Average Condenser Heat Duty'] = df['CONDENSER_HEAT_DUTY'].mean()

    return analysis_results, df, {}

def generate_plots(df):
    """Generates and saves temperature profile, DP, and energy plots."""
    try:
        plt.figure(figsize=(10, 6))
        
        if 'DateAndTime' in df.columns:
            df.sort_values(by='DateAndTime', inplace=True)
            x_axis = df['DateAndTime']
            
            temp_tags = ['TI-13', 'TI-14', 'TI-15', 'TI-16', 'TI-17', 'TI-18', 'TI-19', 'TI-20', 'TI-21', 'TI-22', 'TI-23', 'TI-24', 'TI-25']
            for tag in temp_tags:
                if tag in df.columns:
                    plt.plot(x_axis, df[tag], label=tag, alpha=0.7)

            plt.title("C-02 Column Temperature Profile Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Temperature ({TAG_UNITS['TI-13']})")
            plt.legend(ncol=2)
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_temp_plot_path)
            plt.close()
            print(f"Temperature profile plot saved to {output_temp_plot_path}")
            
    except Exception as e:
        print(f"Error generating temperature plot: {e}")
        
    # Differential Pressure Plot
    try:
        if 'DIFFERENTIAL_PRESSURE' in df.columns:
            plt.figure(figsize=(10, 6))
            plt.plot(df['DateAndTime'], df['DIFFERENTIAL_PRESSURE'], color='purple', alpha=0.8)
            plt.title("C-02 Differential Pressure Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Differential Pressure ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})")
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_dp_plot_path)
            plt.close()
            print(f"Differential pressure plot saved to {output_dp_plot_path}")
    except Exception as e:
        print(f"Error generating DP plot: {e}")

    # Daily Trends Plot
    try:
        df['DATE'] = df['DateAndTime'].dt.date
        daily_trends = df.groupby('DATE').agg({
            'FT-03': 'mean',
            'TI-28': 'mean', 
            'DIFFERENTIAL_PRESSURE': 'mean'
        }).reset_index()

        plt.figure(figsize=(12, 8))
        plt.plot(daily_trends['DATE'], daily_trends['FT-03'], label=f"Avg Top Product Flow ({TAG_UNITS['FT-03']})")
        plt.plot(daily_trends['DATE'], daily_trends['TI-28'], label=f"Avg Top Product Temp ({TAG_UNITS['TI-28']})")
        plt.plot(daily_trends['DATE'], daily_trends['DIFFERENTIAL_PRESSURE'], label=f"Avg DP ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})")
        plt.title("C-02 Daily Trends")
        plt.xlabel("Date")
        plt.ylabel("Value")
        plt.legend()
        plt.grid(True)
        plt.tight_layout()
        plt.savefig(output_trends_plot_path)
        plt.close()
        print(f"Daily trends plot saved to {output_trends_plot_path}")
    except Exception as e:
        print(f"Error generating daily trends plot: {e}")
        
def generate_word_report(analysis_results, df, outliers):
    """Creates a detailed analysis report in a Word document."""
    doc = Document()
    doc.add_heading('C-02 Light Oil Recovery Column Analysis Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    
    summary_text = ""
    if 'Average Reflux Ratio' in analysis_results and isinstance(analysis_results['Average Reflux Ratio'], (float, int)):
        summary_text += f"The column operated with an average reflux ratio of {analysis_results['Average Reflux Ratio']:.2f}, indicating effective control over product separation. "
    
    if 'Overall Material Balance Error (%)' in analysis_results:
        summary_text += f"An overall material balance error of {analysis_results['Overall Material Balance Error (%)']:.2f}% was calculated for the entire process. "
    
    if 'C-02 Naphthalene Loss Status' in analysis_results:
        summary_text += analysis_results['C-02 Naphthalene Loss Status']
        if 'Naphthalene in Top Product (%)' in analysis_results:
            summary_text += f" (Current loss: {analysis_results['Naphthalene in Top Product (%)']:.2f}%)"
    
    doc.add_paragraph(summary_text)

    # Section 2: Key Performance Indicators
    doc.add_heading('2. Key Performance Indicators (KPIs)', level=1)
    doc.add_paragraph("All values are averages over the analysis period.")
    for key, value in analysis_results.items():
        if key == 'Component-wise Material Balance':
            continue
        tag_match = re.search(r'\((.*?)\)', key)
        if tag_match:
            tag = tag_match.group(1)
            unit = TAG_UNITS.get(tag, '')
        else:
            unit = TAG_UNITS.get(key.split(' ')[-1].strip(), '')

        if isinstance(value, str):
            doc.add_paragraph(f"• {key}: {value}")
        else:
            doc.add_paragraph(f"• {key}: {value:.2f} {unit}")

    # Section 3: Component-wise Material Balance
    doc.add_heading('3. Component-wise Material Balance', level=1)
    doc.add_paragraph("This section details the material balance for key components, checking for losses across the system.")
    if 'Component-wise Material Balance' in analysis_results:
        comp_balance = analysis_results['Component-wise Material Balance']
        for comp_key, comp_value in comp_balance.items():
            doc.add_paragraph(f"• {comp_key}: {comp_value:.2f}%")

    # Section 4: Performance Plots
    doc.add_heading('4. Performance Plots', level=1)

    doc.add_heading('4.1 Temperature Profile', level=2)
    doc.add_paragraph("The temperature profile plot shows the gradient across the column.")
    doc.add_picture(output_temp_plot_path, width=Inches(6))

    doc.add_heading('4.2 Differential Pressure (DP)', level=2)
    doc.add_paragraph("Differential pressure is a key indicator of flooding or fouling.")
    doc.add_picture(output_dp_plot_path, width=Inches(6))

    doc.add_heading('4.3 Daily Trends', level=2)
    doc.add_paragraph("This plot shows the daily average trends of key variables.")
    doc.add_picture(output_trends_plot_path, width=Inches(6))

    doc.save(output_report_path)
    print(f"Analysis report generated successfully at {output_report_path}")

def main():
    """Main execution function."""
    engine = connect_to_database()
    if engine is None:
        return

    scada_data = get_scada_data(engine)
    if scada_data is None:
        return

    analysis_results, scada_data, outliers = perform_analysis(scada_data)
    
    if analysis_results:
        generate_plots(scada_data)
        generate_word_report(analysis_results, scada_data, outliers)
        print("C-02 analysis complete.")
    else:
        print("Analysis failed: no data to process.")

if __name__ == "__main__":
    main()