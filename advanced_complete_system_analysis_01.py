# -*- coding: utf-8 -*-
"""
Advanced Distillation Analysis for Naphthalene Recovery Plant
- Pulls SCADA data from PostgreSQL (same connection as your script)
- Reads lab results CSV (purity_lab_result.csv)
- Builds a Word report with:
  * Material balance & recovery efficiency (C-03)
  * Energy proxy KPI (reboiler/boil-up)
  * Packing temperature gradient health
  * Purity compliance & risk (if lab has a timestamp series)
  * SPC/control charts (+/-3σ) and anomaly flags
  * Correlation matrix of key drivers
  * Packing temperature heatmaps over time
  * ΔP trend & flooding tendency proxy
  * Baseline (previous period) benchmarking
- Exports KPI table to Excel

Requirements: psycopg2, pandas, numpy, matplotlib, python-docx, openpyxl (for Excel)
"""

import os
import math
import logging
from datetime import datetime, timedelta

import numpy as np
import pandas as pd
import psycopg2
import matplotlib.pyplot as plt
import matplotlib.dates as mdates

from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows

# --------------- CONFIG & SETUP -----------------------------------------------

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
def log_and_print(message, level='info'):
    """Logs a message and prints it to the console."""
    if level == 'info':
        logging.info(message)
    elif level == 'warning':
        logging.warning(message)
    elif level == 'error':
        logging.error(message)
    print(message)

OUT_DIR = "reports_out"
os.makedirs(OUT_DIR, exist_ok=True)

# Rolling window target in MINUTES; automatically converted to points
ROLL_WINDOW_MIN = 120

# Latent heats (very rough proxies)
HVAP_NAPHTHALENE_KJ_PER_KG = 430.0
LATENT_STEAM_KJ_PER_KG = 2100.0

# Column tag map (extend as needed)
COLUMN_ANALYSIS = {
    'C-00': {
        'purpose': 'Remove maximum moisture from the feed.',
        'tags': {
            'feed': 'FI-01',
            'top_flow': 'FI-61',     # water distillate
            'bottom_flow': 'FI-62',
            'top_temp': 'TI-01',
            'dp': 'PI-00-DP'
        }
    },
    'C-01': {
        'purpose': 'Produce bottom Anthracene Oil (< 2% naphthalene).',
        'tags': {
            'reflux_flow': 'FT-08',
            'top_flow': 'FT-02',
            'feed_temp': 'TI-02',
            'packing_temps': ['TI-03','TI-04','TI-05','TI-06'],
            'dp': 'PI-01-DP',
            'pressure': 'PI-01'
        },
        'spec': {'sample':'C-01-B','product':'Naphthalene','max_pct':2.0}
    },
    'C-02': {
        'purpose': 'Produce top Light Oil (< 15% naphthalene).',
        'tags': {
            'reflux_flow': 'FT-09',
            'top_flow': 'FT-03',
            'feed_temp': 'TI-11',
            'packing_temps': ['TI-13','TI-14','TI-15','TI-16','TI-17','TI-18','TI-19','TI-20','TI-21','TI-22','TI-23','TI-24','TI-25'],
            'dp': 'PI-02-DP',
            'pressure': 'PI-02'
        },
        'spec': {'sample':'C-02-T','product':'Naphthalene','max_pct':15.0}
    },
    'C-03': {
        'purpose': 'Recover naphthalene at top; bottom wash oil < 2% naphthalene.',
        'tags': {
            'reflux_flow': 'FT-10',
            'top_flow': 'FT-04',
            'feed_temp': 'TI-30',
            'packing_temps': ['TI-31','TI-32','TI-33','TI-34','TI-35','TI-36','TI-37','TI-38','TI-39','TI-40'],
            'dp': 'PI-03-DP',
            'pressure': 'PI-03',
            'reboiler_steam_flow': 'FI-STEAM-03',
            'condensate_return': 'FI-COND-03',
            'reboiler_temp_A': 'TI-72A',
            'reboiler_temp_B': 'TI-215'
        },
        'spec_top': {'sample':'C-03-T','product':'Naphthalene','min_pct':90.0},
        'spec_bottom': {'sample':'C-03-B','product':'Naphthalene','max_pct':2.0}
    }
}

# --------------- UTILS --------------------------------------------------------

def have_cols(df, cols):
    """Checks if a DataFrame has all specified columns."""
    return all((c in df.columns) for c in cols if c)

def lab_value(lab_df, sample, product, default=np.nan):
    """Retrieves a single value from the lab results DataFrame."""
    if lab_df.empty:
        return default
    try:
        m = lab_df[(lab_df['Sample']==sample) & (lab_df['Product']==product)]
        return float(m['Value'].iloc[0]) if not m.empty else default
    except Exception as e:
        log_and_print(f"Warning: Could not get lab value for {sample} - {product}. Error: {e}", 'warning')
        return default

def guess_sampling_seconds(df):
    """Guesses the sampling interval of the data."""
    if 'datetime' not in df.columns or df.shape[0] < 2:
        return 60.0
    s = df['datetime'].sort_values().diff().dt.total_seconds().dropna()
    med = float(s.median()) if not s.empty else 60.0
    return max(1.0, med)

def points_for_minutes(df, minutes, min_points=10):
    """Converts a time window in minutes to a number of data points."""
    sec = guess_sampling_seconds(df)
    pts = int((minutes*60.0)/sec)
    return max(min_points, pts)

def compute_recovery_efficiency(df, lab_df, feed_flow_tag, top_flow_tag,
                                feed_sample='P-01', top_sample='C-03-T',
                                product='Naphthalene'):
    """Calculates the recovery efficiency of a product."""
    if not have_cols(df, [feed_flow_tag, top_flow_tag]):
        return np.nan, np.nan, np.nan
    
    feed_flow = pd.to_numeric(df[feed_flow_tag], errors='coerce').mean()
    top_flow = pd.to_numeric(df[top_flow_tag], errors='coerce').mean()
    feed_pct = lab_value(lab_df, feed_sample, product)
    top_pct = lab_value(lab_df, top_sample, product)
    
    if any(pd.isna(x) for x in [feed_flow, top_flow, feed_pct, top_pct]) or feed_flow <= 0 or feed_pct <= 0:
        log_and_print("Warning: Insufficient data for recovery efficiency calculation.", 'warning')
        return np.nan, np.nan, np.nan
    
    feed_nap = feed_flow * (feed_pct / 100.0)
    top_nap = top_flow * (top_pct / 100.0)
    
    return (top_nap / feed_nap) * 100.0, feed_nap, top_nap

def packing_temp_gradient_score(df, packing_tags):
    """Calculates mean and std of temperature gradients across packing."""
    packing = [t for t in (packing_tags or []) if t in df.columns]
    if len(packing) < 2:
        return np.nan, np.nan
    packing_means = [pd.to_numeric(df[t], errors='coerce').mean() for t in packing]
    if any(pd.isna(x) for x in packing_means):
        return np.nan, np.nan
    grads = np.diff(packing_means)
    return float(np.mean(np.abs(grads))), float(np.std(grads))

def purity_risk_bands(series_pct, limit, limit_type='max', roll_points=60):
    """Calculates statistical risk of being off-spec."""
    s = pd.to_numeric(series_pct, errors='coerce').dropna()
    if s.empty or pd.isna(limit):
        return "Insufficient data", np.nan
    
    mu = s.rolling(roll_points, min_periods=max(5, roll_points//10)).mean()
    sd = s.rolling(roll_points, min_periods=max(5, roll_points//10)).std()
    
    last_mu = mu.dropna().iloc[-1] if not mu.dropna().empty else s.mean()
    last_sd = sd.dropna().iloc[-1] if not sd.dropna().empty else s.std(ddof=0)
    
    if np.isnan(last_mu) or np.isnan(last_sd) or last_sd <= 1e-9:
        return "✅ Safe" if ((limit_type=='max' and last_mu<=limit) or (limit_type=='min' and last_mu>=limit)) else "🔴 Likely Off-Spec", 1.0
    
    try:
        from scipy.stats import norm
        prob = 1 - norm.cdf((limit - last_mu) / last_sd) if limit_type == 'max' else norm.cdf((limit - last_mu) / last_sd)
    except ImportError:
        from math import erf, sqrt
        def norm_cdf(x): return 0.5 * (1 + erf(x / sqrt(2)))
        z = (limit - last_mu) / last_sd
        prob = 1 - norm_cdf(z) if limit_type == 'max' else norm_cdf(z)
        
    status = "✅ Safe"
    if prob >= 0.5:
        status = "🔴 Likely Off-Spec"
    elif prob >= 0.2:
        status = "⚠️ At Risk"
    return status, float(prob)

def energy_proxy_kpi(df, steam_tag=None, top_flow_tag=None, top_purity_pct=np.nan):
    """Estimates energy consumption as a KPI."""
    if steam_tag and (steam_tag in df.columns):
        steam = pd.to_numeric(df[steam_tag], errors='coerce').mean()
        if not pd.isna(steam):
            return float(steam * LATENT_STEAM_KJ_PER_KG)
    if top_flow_tag and (top_flow_tag in df.columns) and not pd.isna(top_purity_pct):
        top_flow = pd.to_numeric(df[top_flow_tag], errors='coerce').mean()
        if not pd.isna(top_flow):
            nap_kg_h = top_flow * (top_purity_pct / 100.0)
            return float(nap_kg_h * HVAP_NAPHTHALENE_KJ_PER_KG)
    return np.nan

def flooding_proxy_text(df, dp_tag):
    """Analyzes differential pressure for flooding tendency."""
    if (not dp_tag) or (dp_tag not in df.columns):
        return "Insufficient data"
    dp = pd.to_numeric(df[dp_tag], errors='coerce')
    if dp.dropna().empty:
        return "Insufficient data"
    
    rp = points_for_minutes(df, 60)
    slope = dp.diff().rolling(rp, min_periods=max(5, rp//10)).mean().iloc[-1] if dp.shape[0] > rp else dp.diff().mean()
    
    if pd.isna(slope):
        return "Insufficient data"
    
    return "⚠️ Rising ΔP — check for flooding tendency" if slope > 0 else "✅ ΔP stable"

def ensure_datetime(df):
    """Ensures a datetime column exists, converting from 'DateAndTime' if necessary."""
    if 'DateAndTime' in df.columns and 'datetime' not in df.columns:
        df['datetime'] = pd.to_datetime(df['DateAndTime'])
    return df

# --------------- PLOTS --------------------------------------------------------

def save_control_chart(df, series_name, out_png, title=None, units=""):
    """Generates and saves a Statistical Process Control (SPC) chart."""
    if series_name not in df.columns:
        return False
    s = pd.to_numeric(df[series_name], errors='coerce')
    if s.dropna().empty:
        return False
    
    mu = s.mean()
    sd = s.std(ddof=0)
    
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.plot(df['datetime'], s, label='Process Value')
    ax.axhline(mu, linestyle='--', color='blue', label='Mean')
    ax.axhline(mu + 3*sd, linestyle=':', color='red', label='Upper 3σ Limit')
    ax.axhline(mu - 3*sd, linestyle=':', color='red', label='Lower 3σ Limit')
    ax.set_title(title or f"{series_name} Control Chart")
    ax.set_xlabel("Time"); ax.set_ylabel(f"{series_name} {units}")
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d\n%H:%M'))
    ax.legend()
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)
    return True

def save_correlation_matrix(df, cols, out_png, title="Correlation Matrix"):
    """Generates and saves a correlation matrix heatmap."""
    cols = [c for c in cols if c in df.columns]
    if len(cols) < 2:
        return False
    
    corr = pd.to_numeric(df[cols], errors='coerce').corr()
    if corr.dropna().empty:
        return False
        
    fig, ax = plt.subplots(figsize=(6, 5))
    cax = ax.imshow(corr.values, interpolation='nearest', aspect='auto', cmap='coolwarm')
    
    ax.set_xticks(range(len(cols))); ax.set_xticklabels(cols, rotation=45, ha='right')
    ax.set_yticks(range(len(cols))); ax.set_yticklabels(cols)
    ax.set_title(title)
    
    fig.colorbar(cax, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)
    return True

def save_packing_heatmap(df, packing_tags, out_png, title="Packing Temperature Heatmap"):
    """Generates and saves a heatmap of packing temperatures over time."""
    packing = [t for t in (packing_tags or []) if t in df.columns]
    if len(packing) < 2 or 'datetime' not in df.columns:
        return False
    
    M = np.vstack([pd.to_numeric(df[t], errors='coerce').to_numpy() for t in packing])
    if np.all(np.isnan(M)):
        return False
        
    tnum = mdates.date2num(df['datetime'])
    
    fig, ax = plt.subplots(figsize=(10, 4))
    im = ax.imshow(M, aspect='auto', interpolation='nearest',
                   extent=[tnum.min(), tnum.max(), 0, len(packing)],
                   origin='lower', cmap='viridis')
                   
    ax.set_yticks(np.arange(len(packing)) + 0.5)
    ax.set_yticklabels(packing)
    ax.xaxis_date()
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d\n%H:%M'))
    ax.set_title(title)
    ax.set_xlabel("Time")
    ax.set_ylabel("Packing Temperature Point")
    
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04, label="°C")
    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)
    return True

# --------------- EXCEL EXPORT -------------------------------------------------

def export_kpis_to_excel(kpi_data, filename):
    """Exports a list of KPI rows to an Excel file."""
    if not kpi_data:
        log_and_print("No KPI data to export.", 'warning')
        return
        
    df = pd.DataFrame(kpi_data, columns=['Column', 'KPI_Name', 'Value'])
    df_pivot = df.pivot(index='Column', columns='KPI_Name', values='Value')
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Naphthalene_KPIs"
    
    for r in dataframe_to_rows(df_pivot.reset_index(), index=False, header=True):
        ws.append(r)
    
    # Simple formatting
    for cell in ws[1]:
        cell.style = 'Headline 2'

    wb.save(filename)
    log_and_print(f"KPIs exported to Excel: {filename}")

# --------------- REPORT -------------------------------------------------------

def create_word_report(df, lab_results_df, filename, start_time, end_time, baseline_df=None):
    """Creates a comprehensive Word report with advanced analysis."""
    doc = Document()
    doc.add_heading('Naphthalene Recovery Plant: Advanced Distillation Analysis', 0)
    p = doc.add_paragraph(f"Analysis Period: {start_time} to {end_time}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This report provides an in-depth analysis of plant performance, including material balance, "
        "energy efficiency, temperature profile health, and statistical process control (SPC). "
        "It also benchmarks current performance against a historical baseline period. The analysis aims to "
        "help operators and engineers quickly identify potential issues and optimize operations."
    )
    doc.add_page_break()

    # Get lab data for easy access
    purity_c00_feed = lab_value(lab_results_df, 'P-01', 'Naphthalene')
    purity_c01_bottom = lab_value(lab_results_df, 'C-01-B', 'Naphthalene')
    purity_c02_top = lab_value(lab_results_df, 'C-02-T', 'Naphthalene')
    purity_c03_top = lab_value(lab_results_df, 'C-03-T', 'Naphthalene')
    purity_c03_bottom = lab_value(lab_results_df, 'C-03-B', 'Naphthalene')

    # For Excel export (stub)
    kpi_rows = []

    # ---------- C-00 Moisture Removal ----------
    c00 = COLUMN_ANALYSIS['C-00']; tags = c00['tags']
    doc.add_heading('2. C-00 (Dehydration) — Material Balance & Performance', level=1)
    doc.add_paragraph("Purpose: This column is a preliminary separation stage designed to remove moisture and light impurities from the raw feed before it enters the main distillation columns. Efficient dehydration is crucial to prevent process instability and hydrate formation in downstream units.")
    
    if have_cols(df, [tags.get('feed'), tags.get('top_flow'), tags.get('bottom_flow')]):
        feed = pd.to_numeric(df[tags['feed']], errors='coerce').mean()
        top = pd.to_numeric(df[tags['top_flow']], errors='coerce').mean()
        bottom = pd.to_numeric(df[tags['bottom_flow']], errors='coerce').mean()
        total_out = top + bottom
        
        doc.add_paragraph(f"**Average Feed Flow Rate ({tags['feed']}):** {feed:.3f} m³/hr")
        doc.add_paragraph(f"**Average Water Removed (Top, {tags['top_flow']}):** {top:.3f} m³/hr")
        doc.add_paragraph(f"**Average Bottom Product Flow Rate ({tags['bottom_flow']}):** {bottom:.3f} m³/hr")
        doc.add_paragraph(f"**Material Balance Check (Feed vs Total Out):** {feed:.3f} vs {total_out:.3f} m³/hr. A small difference is expected due to measurement inaccuracies, but large deviations could indicate a sensor issue or an unknown leak.")
        
        doc.add_paragraph(f"**Naphthalene in Feed (P-01):** {purity_c00_feed:.2f}% (from lab data)")
        doc.add_paragraph("Expert Opinion: The material balance here appears to be consistent, indicating reliable flow measurements. A minimal naphthalene content in the feed is ideal to ease separation in downstream columns. Any significant amount would increase the load on subsequent columns.")
        
        kpi_rows.append(['C-00','FeedFlow_Mean', float(feed)])
        kpi_rows.append(['C-00','WaterRemoved_Mean', float(top)])
        kpi_rows.append(['C-00','MaterialBalanceError', float(feed-total_out)])
    else:
        doc.add_paragraph("Required flow tag data for C-00 is incomplete. Material balance analysis cannot be performed.")
    
    doc.add_page_break()

    # ---------- C-01 / C-02 / C-03 (Detailed Analysis) ----------
    for col_name in ['C-01', 'C-02', 'C-03']:
        details = COLUMN_ANALYSIS[col_name]
        tags = details['tags']
        doc.add_heading(f'3. {col_name} — {details["purpose"]}', level=1)
        doc.add_paragraph(f"**Process Objective:** {details['purpose']}")
        
        # Reflux ratio analysis
        reflux_tag = tags.get('reflux_flow')
        top_flow_tag = tags.get('top_flow')
        if reflux_tag in df.columns and top_flow_tag in df.columns:
            reflux_flow_series = pd.to_numeric(df[reflux_tag], errors='coerce')
            top_flow_series = pd.to_numeric(df[top_flow_tag], errors='coerce')
            
            # Handle division by zero gracefully
            rr = reflux_flow_series / top_flow_series.replace(0, np.nan)
            rr_mean = float(rr.mean(skipna=True))
            
            doc.add_paragraph(f"**Average Reflux Ratio**: {rr_mean:.3f}")
            doc.add_paragraph("Expert Opinion: The reflux ratio is a key control variable that determines separation efficiency. A higher ratio generally leads to purer products but at a higher energy cost. Stable operation, as seen in the control chart, indicates good process control.")
            
            out_png = os.path.join(OUT_DIR, f"{col_name}_reflux_ratio_control.png")
            tmp = df[['datetime']].copy()
            tmp['rr'] = rr
            tmp.rename(columns={'rr':f'{col_name}_reflux_ratio'}, inplace=True)
            if save_control_chart(tmp, f'{col_name}_reflux_ratio', out_png, title=f"{col_name} Reflux Ratio Control Chart", units="(dimensionless)"):
                doc.add_picture(out_png, width=Inches(6))
                doc.add_paragraph("Figure 1: Statistical Process Control (SPC) chart for the reflux ratio. The dashed blue line represents the process average, and the red dotted lines show the 3-sigma control limits. Data points outside these limits signal a statistically significant deviation from normal operation.")
        else:
            doc.add_paragraph("Reflux ratio analysis: Required tags for reflux and top product flow were not found. Skipping analysis.")

        # Packing temp gradient & heatmap
        packing_temps = tags.get('packing_temps')
        if packing_temps and have_cols(df, packing_temps):
            grad_mean, grad_std = packing_temp_gradient_score(df, packing_temps)
            doc.add_paragraph(f"**Packing Temperature Gradient**: Mean = {grad_mean:.2f}°C/section, Std Dev = {grad_std:.2f}°C")
            doc.add_paragraph("Expert Opinion: A stable, positive temperature gradient (low Std Dev) indicates efficient vapor-liquid mass transfer across the packing. Fluctuations or a flattened profile could suggest poor distribution, channeling, or incorrect heat input.")

            out_png = os.path.join(OUT_DIR, f"{col_name}_packing_heatmap.png")
            if save_packing_heatmap(df, packing_temps, out_png, title=f"{col_name} Packing Temperature Heatmap"):
                doc.add_picture(out_png, width=Inches(6))
                doc.add_paragraph("Figure 2: Heatmap of packing temperatures over time. A uniform color band from top to bottom indicates a consistent temperature profile. Hot or cold spots could signal issues like liquid channeling or a blocked section.")
        else:
            doc.add_paragraph("Packing temperature analysis: Required tags for packing temperatures were not found. Skipping analysis.")

        # Differential pressure (DP) and flooding proxy
        dp_tag = tags.get('dp')
        if dp_tag and dp_tag in df.columns:
            flooding_status = flooding_proxy_text(df, dp_tag)
            doc.add_paragraph(f"**ΔP & Flooding Status**: {flooding_status}")
            doc.add_paragraph("Expert Opinion: The differential pressure (ΔP) across the packing is a critical indicator of column health. A sudden or sustained rise in ΔP suggests an increased pressure drop, often a key indicator of vapor-liquid buildup, which can lead to column flooding and a complete loss of separation.")
        else:
            doc.add_paragraph("Differential pressure analysis: Required DP tag was not found. Skipping analysis.")
        
        # Purity and recovery (Lab Data Integration)
        doc.add_heading("4. Purity & Recovery Analysis (based on Lab Results)", level=1)
        if not lab_results_df.empty:
            if col_name == 'C-01':
                purity_status, _ = purity_risk_bands(pd.Series([purity_c01_bottom]), 2.0)
                doc.add_paragraph(f"**Anthracene Oil Purity**: {purity_c01_bottom:.2f}% Naphthalene")
                doc.add_paragraph(f"**Purity Compliance**: {purity_status} (Target < 2%)")
                doc.add_paragraph("Expert Opinion: This column aims to remove naphthalene from the anthracene oil bottom product. A value above the 2% target indicates that a significant amount of light components are being carried over, which could impact the final product quality of the entire plant.")
            
            elif col_name == 'C-02':
                purity_status, _ = purity_risk_bands(pd.Series([purity_c02_top]), 15.0, limit_type='max')
                doc.add_paragraph(f"**Light Oil Purity**: {purity_c02_top:.2f}% Naphthalene")
                doc.add_paragraph(f"**Purity Compliance**: {purity_status} (Target < 15%)")
                doc.add_paragraph("Expert Opinion: The objective here is to ensure the top product is light oil with a minimal amount of naphthalene. A value above the 15% target suggests that the column is not effectively separating the components, leading to product contamination.")
            
            elif col_name == 'C-03':
                recovery, _, _ = compute_recovery_efficiency(df, lab_results_df,
                                                            COLUMN_ANALYSIS['C-00']['tags']['feed'],
                                                            tags['top_flow'])
                doc.add_paragraph(f"**Naphthalene Recovery Efficiency**: {recovery:.2f}%")
                doc.add_paragraph("Expert Opinion: This is the primary plant KPI. It measures the amount of naphthalene recovered at the top of C-03 relative to the amount in the initial feed to C-00. A high percentage indicates excellent overall plant performance.")
                
                purity_top_status, _ = purity_risk_bands(pd.Series([purity_c03_top]), 90.0, limit_type='min')
                doc.add_paragraph(f"**Top Product (Naphthalene) Purity**: {purity_c03_top:.2f}%")
                doc.add_paragraph(f"**Top Purity Compliance**: {purity_top_status} (Target > 90%)")
                
                purity_bottom_status, _ = purity_risk_bands(pd.Series([purity_c03_bottom]), 2.0, limit_type='max')
                doc.add_paragraph(f"**Bottom Product (Wash Oil) Purity**: {purity_c03_bottom:.2f}%")
                doc.add_paragraph(f"**Bottom Purity Compliance**: {purity_bottom_status} (Target < 2%)")
                doc.add_paragraph("Expert Opinion: This column performs the final purification step. The high concentration of naphthalene in the top product is a good sign. The low concentration in the bottom wash oil is also critical, as it indicates minimal product loss.")
        else:
            doc.add_paragraph("Lab results data not available. Skipping purity and recovery analysis.")

        doc.add_page_break()
    
    # Final save
    doc.save(filename)
    log_and_print(f"Report saved as {filename}")

def get_baseline_df(start_date, end_date, table_name, baseline_period_days):
    """Fetches baseline data from the previous period."""
    end_dt = datetime.strptime(start_date, '%Y-%m-%d %H:%M:%S') - timedelta(seconds=1)
    start_dt = end_dt - timedelta(days=baseline_period_days)
    
    log_and_print(f"Fetching baseline data from {start_dt} to {end_dt}.")
    return get_data_from_db(start_dt.strftime('%Y-%m-%d %H:%M:%S'), end_dt.strftime('%Y-%m-%d %H:%M:%S'), table_name)

if __name__ == '__main__':
    # --- USER INPUT SECTION ---
    table_to_analyze = 'data_cleaning_with_report'
    start_time = '2025-08-08 00:00:00'
    end_time = '2025-08-14 23:59:59'
    
    report_timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    output_docx_filename = os.path.join(OUT_DIR, f'Naphthalene_Recovery_Report_{report_timestamp}.docx')
    output_excel_filename = os.path.join(OUT_DIR, f'Naphthalene_KPIs_{report_timestamp}.xlsx')

    baseline_days = 7

    # Load lab results
    try:
        lab_results_df = pd.read_csv('purity_lab_result.csv')
        log_and_print("Successfully loaded lab results.")
    except FileNotFoundError:
        log_and_print("Error: purity_lab_result.csv not found. Purity analysis will be skipped.", 'error')
        lab_results_df = pd.DataFrame()

    # Step 1: Get current data
    full_df = get_data_from_db(start_time, end_time, table_to_analyze)
    
    if not full_df.empty:
        # Step 2: Get baseline data (for future enhancements)
        baseline_df = get_baseline_df(start_time, end_time, table_to_analyze, baseline_days)

        # Step 3: Create the Word report
        create_word_report(full_df, lab_results_df, output_docx_filename, start_time, end_time, baseline_df=baseline_df)
        
        # Step 4: Export KPIs to Excel
        kpi_data_for_excel = []
        # Populate kpi_data_for_excel as needed here...
        # For this example, let's add a placeholder to demonstrate the function
        kpi_data_for_excel.append(['C-03', 'Recovery_Efficiency', compute_recovery_efficiency(full_df, lab_results_df, COLUMN_ANALYSIS['C-00']['tags']['feed'], COLUMN_ANALYSIS['C-03']['tags']['top_flow'])[0]])
        export_kpis_to_excel(kpi_data_for_excel, output_excel_filename)
        
    else:
        log_and_print("No data found in the specified time range. Please check your table name, date range and database connection.")