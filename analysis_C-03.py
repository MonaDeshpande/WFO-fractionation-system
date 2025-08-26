import pandas as pd
from sqlalchemy import create_engine
import psycopg2
from datetime import datetime
import matplotlib.pyplot as plt
import os
import numpy as np
from docx import Document
from docx.shared import Inches
import sqlalchemy
import re

# Database connection parameters (update with your actual details)
DB_HOST = "localhost"
DB_NAME = "scada_data_analysis"
DB_USER = "postgres"
DB_PASSWORD = "ADMIN"

# Define units for each tag
TAG_UNITS = {
    'FT-06': 'kg/h', # Feed flow from C-02 bottom
    'TI-30': 'degC', # Feed temperature
    'TI-31': 'degC', # Column temperature
    'TI-32': 'degC', # Column temperature
    'TI-33': 'degC', # Column temperature
    'TI-34': 'degC', # Column temperature
    'TI-35': 'degC', # Column temperature
    'TI-36': 'degC', # Column temperature
    'TI-37': 'degC', # Column temperature
    'TI-38': 'degC', # Column temperature
    'TI-39': 'degC', # Column temperature
    'TI-40': 'degC', # Column temperature (vent line)
    'TI-41': 'degC', # After main condenser
    'TI-42': 'degC', # After vent condenser
    'TI-43': 'degC', # Liquid from vent condenser
    'FT-10': 'kg/h', # Reflux flow rate
    'FT-04': 'kg/h', # Top product flow rate
    'TI-44': 'degC', # Top product temperature (after cooler)
    'LI-05': '%',    # Column bottom level
    'FT-07': 'kg/h', # Bottom product flow rate
    'FT-05': 'kg/h', # This tag is from C-01, but in the old C-03 code it was used. FT-07 is the correct tag for C-03 bottoms.
    'TI-45': 'degC', # Bottom product temperature (after heater)
    'TI-73A': 'degC', # Reboiler thermic fluid inlet temp
    'TI-73B': 'degC', # Reboiler thermic fluid outlet temp
    'PTB-03': 'mmHg', # Column bottom pressure
    'PTT-03': 'mmHg', # Column top pressure
    'FT-104': 'm3/h', # Cooling water flow to vent condenser
    'FT-202': 'kg/h', # C-03 Reboiler thermic Fluid flow rate
    'DIFFERENTIAL_PRESSURE': 'mmHg',
    'REBOILER_HEAT_DUTY': 'kW',
    'CONDENSER_HEAT_DUTY': 'kW',
    'REFLUX_RATIO': '',
    'MATERIAL_BALANCE_ERROR': '%',
    'NAPHTHALENE_PURITY_PERCENTAGE': '%',
    'THIANAPHTHENE_PERCENTAGE': '%',
    'QUINOLINE_PPM': 'ppm',
    'UNKNOWN_IMPURITY_PERCENTAGE': '%'
}

# File paths for saving generated plots and report
output_report_path = "C-03_Analysis_Report.docx"
output_temp_plot_path = "C-03_Temperature_Profile.png"
output_dp_plot_path = "C-03_Differential_Pressure.png"
output_trends_plot_path = "C-03_Daily_Trends.png"

# Engineering constants for heat duty calculations
THERMIC_FLUID_SPECIFIC_HEAT = 2.0  # kJ/(kg·°C) - Assumed value
WATER_SPECIFIC_HEAT = 4.186       # kJ/(kg·°C)

# File path for composition data. The code now simulates reading this.
COMPOSITION_FILE_PATH = "your_composition_data.csv"

def connect_to_database():
    """Establishes a connection to the PostgreSQL database."""
    try:
        engine = create_engine(f'postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}/{DB_NAME}')
        print("Database connection successful.")
        return engine
    except Exception as e:
        print(f"Error connecting to the database: {e}")
        return None

def get_scada_data(engine):
    """Retrieves specific SCADA data for the C-03 column."""
    try:
        desired_columns = [
            "DateAndTime", "FT-06", "TI-30", "TI-31", "TI-32", "TI-33", "TI-34", "TI-35", "TI-36", "TI-37",
            "TI-38", "TI-39", "TI-40", "TI-41", "TI-42", "TI-43", "FT-10", "FT-04", "TI-44", "LI-05",
            "FT-07", "TI-45", "TI-73A", "TI-73B", "PTB-03", "PTT-03", "FT-104", "FT-202"
        ]
        
        inspector = sqlalchemy.inspect(engine)
        columns = inspector.get_columns('wide_scada_data')
        column_names = [col['name'] for col in columns]
        
        final_columns = []
        for d_col in desired_columns:
            for db_col in column_names:
                if d_col.lower().replace('-', '') == db_col.lower().replace('-', ''):
                    final_columns.append(f'"{db_col}"')
                    break
        
        if not final_columns:
            print("Error: No matching columns found for C-03. Data retrieval failed.")
            return None

        select_clause = ", ".join(final_columns)
        query = f"""
        SELECT {select_clause}
        FROM wide_scada_data
        WHERE "DateAndTime" BETWEEN '2025-08-08 00:00:00' AND '2025-08-20 23:59:59'
        ORDER BY "DateAndTime";
        """
        
        df = pd.read_sql(query, engine)
        df['DateAndTime'] = pd.to_datetime(df['DateAndTime'])
        print("SCADA data for C-03 retrieved successfully.")
        return df
    except Exception as e:
        print(f"Error retrieving SCADA data: {e}")
        return None

def get_composition_data(file_path):
    """
    Simulates reading composition data from a CSV file.
    It now simulates reading from your provided image.
    """
    try:
        # Simulate reading the data from the image you provided.
        composition_data = {
            'Analysis Date': ['08.08.25', '08.08.25', '08.08.25', '08.08.25', '08.08.25', '08.08.25', '08.08.25', '08.08.25', '08.08.25', '08.08.25'],
            'Analysis Time': ['09.15AM', '06.00PM', '06.00PM', '06.00PM', '06.00PM', '06.00PM', '06.00AM', '06.00AM', '06.00AM', '09.30AM'],
            'Sample Detail': ['P-01', 'P-01', 'C-03-T', 'C-02-T', 'C-03-B', 'C-01-B', 'C-03-T', 'C-02-T', 'C-03-B', 'P-01'],
            'Material': ['WFO', 'WFO', 'NO', 'LCO', 'WO-270C', 'ATO', 'NO', 'LCO', 'WO-270C', 'WFO'],
            'Naphth. % by GC': [56.53, 55.52, 88.94, 19.74, 1.46, 0.01, 87.71, 53.42, 1.25, 56.38],
            'Thianaphth. %': [2.04, 2.02, 4.22, np.nan, np.nan, np.nan, 4.37, np.nan, np.nan, 2.03],
            'Quinoline in ppm': [17459, 17442, 6502, np.nan, np.nan, np.nan, 23446, np.nan, np.nan, 18189],
            'Unknown Impurity%': [1.83, 1.84, 1.7, np.nan, np.nan, np.nan, 0.75, np.nan, np.nan, 1.9],
        }
        composition_df = pd.DataFrame(composition_data)
        
        # Get feed composition from C-03 feed (which is C-02 bottom product from the image)
        feed_composition = {}
        c03_feed = composition_df[composition_df['Sample Detail'] == 'C-02-T'] # C-02 bottom is not available, so we use C-02 top as a placeholder for the feed. This is an assumption.
        if not c03_feed.empty:
            feed_composition['Naphth. % by GC'] = c03_feed['Naphth. % by GC'].iloc[-1]
            feed_composition['Thianaphth. %'] = c03_feed['Thianaphth. %'].iloc[-1] if 'Thianaphth. %' in c03_feed.columns else np.nan
            feed_composition['Quinoline in ppm'] = c03_feed['Quinoline in ppm'].iloc[-1] if 'Quinoline in ppm' in c03_feed.columns else np.nan
            feed_composition['Unknown Impurity%'] = c03_feed['Unknown Impurity%'].iloc[-1] if 'Unknown Impurity%' in c03_feed.columns else np.nan

        # Get top product composition from C-03 top product from the image
        top_product_composition = {}
        c03_top = composition_df[composition_df['Sample Detail'] == 'C-03-T']
        if not c03_top.empty:
            top_product_composition['Naphth. % by GC'] = c03_top['Naphth. % by GC'].iloc[-1]
            top_product_composition['Thianaphth. %'] = c03_top['Thianaphth. %'].iloc[-1] if 'Thianaphth. %' in c03_top.columns else np.nan
            top_product_composition['Quinoline in ppm'] = c03_top['Quinoline in ppm'].iloc[-1] if 'Quinoline in ppm' in c03_top.columns else np.nan
            top_product_composition['Unknown Impurity%'] = c03_top['Unknown Impurity%'].iloc[-1] if 'Unknown Impurity%' in c03_top.columns else np.nan
        
        # Get bottom product composition from C-03 bottom product (from the image)
        bottom_product_composition = {}
        c03_bottom = composition_df[composition_df['Sample Detail'] == 'C-03-B']
        if not c03_bottom.empty:
            bottom_product_composition['Naphth. % by GC'] = c03_bottom['Naphth. % by GC'].iloc[-1]
            bottom_product_composition['Thianaphth. %'] = c03_bottom['Thianaphth. %'].iloc[-1] if 'Thianaphth. %' in c03_bottom.columns else np.nan
            bottom_product_composition['Quinoline in ppm'] = c03_bottom['Quinoline in ppm'].iloc[-1] if 'Quinoline in ppm' in c03_bottom.columns else np.nan
            bottom_product_composition['Unknown Impurity%'] = c03_bottom['Unknown Impurity%'].iloc[-1] if 'Unknown Impurity%' in c03_bottom.columns else np.nan

        # Replace NaN with a default value (e.g., 0) for calculations
        for comp in [feed_composition, top_product_composition, bottom_product_composition]:
            for key, value in comp.items():
                if pd.isna(value):
                    comp[key] = 0.0

        return feed_composition, top_product_composition, bottom_product_composition
    
    except FileNotFoundError:
        print(f"Warning: Composition data file not found at {file_path}. Using default.")
        return {}, {}, {}
    except Exception as e:
        print(f"Error reading composition data: {e}. Using default.")
        return {}, {}, {}

def perform_analysis(df):
    """
    Performs key calculations for C-03, including material/energy balances
    and reflux ratio.
    """
    if df is None or df.empty:
        return {}, df, {}

    outliers = {}
    analysis_results = {}
    
    # Material Balance Analysis
    if all(tag in df.columns for tag in ['FT-06', 'FT-04', 'FT-07']):
        feed_flow_avg = df['FT-06'].mean()
        top_product_flow_avg = df['FT-04'].mean()
        bottom_product_flow_avg = df['FT-07'].mean()
        
        analysis_results['Average Feed Flow (FT-06)'] = feed_flow_avg
        analysis_results['Average Top Product Flow (FT-04)'] = top_product_flow_avg
        analysis_results['Average Bottom Product Flow (FT-07)'] = bottom_product_flow_avg
        
        if feed_flow_avg > 0:
            material_balance_error = ((feed_flow_avg - (top_product_flow_avg + bottom_product_flow_avg)) / feed_flow_avg) * 100
            analysis_results['Material Balance Error (%)'] = abs(material_balance_error)

    # Naphthalene Purity & Impurity Analysis
    feed_comp, top_prod_comp, bottom_prod_comp = get_composition_data(COMPOSITION_FILE_PATH)
    
    # Calculate Naphthalene Loss in C-03 Bottoms
    naphthalene_loss_pct = "N/A"
    if 'Naphth. % by GC' in feed_comp and 'Naphth. % by GC' in bottom_prod_comp and feed_flow_avg > 0 and bottom_product_flow_avg > 0:
        feed_naphthalene_mass = feed_flow_avg * (feed_comp['Naphth. % by GC'] / 100)
        bottoms_naphthalene_mass = bottom_product_flow_avg * (bottom_prod_comp['Naphth. % by GC'] / 100)
        
        if feed_naphthalene_mass > 0:
            naphthalene_loss_pct = (bottoms_naphthalene_mass / feed_naphthalene_mass) * 100
            analysis_results['Naphthalene Loss in C-03 Bottoms (%)'] = naphthalene_loss_pct
    else:
        analysis_results['Naphthalene Loss in C-03 Bottoms (%)'] = "N/A (Missing data)"

    if top_prod_comp and 'Naphth. % by GC' in top_prod_comp:
        analysis_results['Naphthalene in Top Product (%)'] = top_prod_comp['Naphth. % by GC']
        analysis_results['Thianaphth. in Top Product (%)'] = top_prod_comp.get('Thianaphth. %', 'N/A')
        analysis_results['Quinoline in Top Product (ppm)'] = top_prod_comp.get('Quinoline in ppm', 'N/A')
        analysis_results['Unknown Impurity in Top Product (%)'] = top_prod_comp.get('Unknown Impurity%', 'N/A')
        
        if top_prod_comp['Naphth. % by GC'] >= 96.0: # Updated purity threshold to 96%
            analysis_results['Naphthalene Purity Status'] = "SUCCESS: Naphthalene purity is at or above target (96%)."
        else:
            analysis_results['Naphthalene Purity Status'] = "ALERT: Naphthalene purity is below target (96%)."

    if bottom_prod_comp and 'Naphth. % by GC' in bottom_prod_comp:
        analysis_results['Naphthalene in Bottom Product (%)'] = bottom_prod_comp['Naphth. % by GC']
    
    # Reflux Ratio
    if 'FT-10' in df.columns and 'FT-04' in df.columns:
        reflux_flow_avg = df['FT-10'].mean()
        top_product_flow_avg = df['FT-04'].mean()
        
        if top_product_flow_avg > 0:
            reflux_ratio = reflux_flow_avg / top_product_flow_avg
            analysis_results['Average Reflux Ratio'] = reflux_ratio
        else:
            analysis_results['Average Reflux Ratio'] = "N/A (Zero Top Product Flow)"
            
    # Differential Pressure (DP) Calculation
    if 'PTT-03' in df.columns and 'PTB-03' in df.columns:
        df['DIFFERENTIAL_PRESSURE'] = df['PTB-03'] - df['PTT-03']
        analysis_results['Average Differential Pressure'] = df['DIFFERENTIAL_PRESSURE'].mean()
        analysis_results['Maximum Differential Pressure'] = df['DIFFERENTIAL_PRESSURE'].max()
        
    # Energy Balance
    # Reboiler Heat Duty
    if all(tag in df.columns for tag in ['TI-73A', 'TI-73B', 'FT-202']):
        df['REBOILER_HEAT_DUTY'] = df['FT-202'] * THERMIC_FLUID_SPECIFIC_HEAT * (df['TI-73B'] - df['TI-73A'])
        analysis_results['Average Reboiler Heat Duty'] = df['REBOILER_HEAT_DUTY'].mean()
    # Using TI-211/TI-212 as alternative reboiler temps
    elif all(tag in df.columns for tag in ['TI-211', 'TI-212', 'FT-202']):
        df['REBOILER_HEAT_DUTY'] = df['FT-202'] * THERMIC_FLUID_SPECIFIC_HEAT * (df['TI-212'] - df['TI-211'])
        analysis_results['Average Reboiler Heat Duty'] = df['REBOILER_HEAT_DUTY'].mean()
    else:
        analysis_results['Average Reboiler Heat Duty'] = 'N/A (Missing data)'

    # Condenser Heat Duty
    if all(tag in df.columns for tag in ['TI-41', 'TI-42', 'FT-104']):
        df['CONDENSER_HEAT_DUTY'] = df['FT-104'] * WATER_SPECIFIC_HEAT * (df['TI-41'] - df['TI-42'])
        analysis_results['Average Condenser Heat Duty'] = df['CONDENSER_HEAT_DUTY'].mean()
    else:
        analysis_results['Average Condenser Heat Duty'] = 'N/A (Missing data)'

    return analysis_results, df, outliers

def generate_plots(df):
    """Generates and saves temperature profile, DP, and daily trends plots."""
    # Temperature Profile Plot
    try:
        plt.figure(figsize=(10, 6))
        
        if 'DateAndTime' in df.columns:
            df.sort_values(by='DateAndTime', inplace=True)
            x_axis = df['DateAndTime']
            
            temp_tags = ['TI-31', 'TI-32', 'TI-33', 'TI-34', 'TI-35', 'TI-36', 'TI-37', 'TI-38', 'TI-39', 'TI-40']
            for tag in temp_tags:
                if tag in df.columns:
                    plt.plot(x_axis, df[tag], label=tag, alpha=0.7)

            plt.title("C-03 Column Temperature Profile Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Temperature ({TAG_UNITS['TI-31']})")
            plt.legend(ncol=2)
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_temp_plot_path)
            plt.close()
            print(f"Temperature profile plot saved to {output_temp_plot_path}")
            
    except Exception as e:
        print(f"Error generating temperature plot: {e}")
        
    # Differential Pressure Plot
    try:
        if 'DIFFERENTIAL_PRESSURE' in df.columns:
            plt.figure(figsize=(10, 6))
            plt.plot(df['DateAndTime'], df['DIFFERENTIAL_PRESSURE'], color='purple', alpha=0.8)
            plt.title("C-03 Differential Pressure Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Differential Pressure ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})")
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_dp_plot_path)
            plt.close()
            print(f"Differential pressure plot saved to {output_dp_plot_path}")
    except Exception as e:
        print(f"Error generating DP plot: {e}")

    # Daily Trends Plot
    try:
        df['DATE'] = df['DateAndTime'].dt.date
        daily_trends = df.groupby('DATE').agg({
            'FT-04': 'mean',
            'TI-44': 'mean', 
            'DIFFERENTIAL_PRESSURE': 'mean'
        }).reset_index()

        plt.figure(figsize=(12, 8))
        plt.plot(daily_trends['DATE'], daily_trends['FT-04'], label=f"Avg Top Product Flow ({TAG_UNITS['FT-04']})")
        plt.plot(daily_trends['DATE'], daily_trends['TI-44'], label=f"Avg Top Product Temp ({TAG_UNITS['TI-44']})")
        plt.plot(daily_trends['DATE'], daily_trends['DIFFERENTIAL_PRESSURE'], label=f"Avg DP ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})")
        plt.title("C-03 Daily Trends")
        plt.xlabel("Date")
        plt.ylabel("Value")
        plt.legend()
        plt.grid(True)
        plt.tight_layout()
        plt.savefig(output_trends_plot_path)
        plt.close()
        print(f"Daily trends plot saved to {output_trends_plot_path}")
    except Exception as e:
        print(f"Error generating daily trends plot: {e}")
        
def generate_word_report(analysis_results, df, outliers):
    """Creates a detailed analysis report in a Word document."""
    doc = Document()
    doc.add_heading('C-03 Naphthalene Oil Column Analysis Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    
    summary_text = ""
    if 'Average Reflux Ratio' in analysis_results and isinstance(analysis_results['Average Reflux Ratio'], (float, int)):
        summary_text += f"The column operated with an average reflux ratio of {analysis_results['Average Reflux Ratio']:.2f}, indicating effective control over product separation. "
    
    if 'Material Balance Error (%)' in analysis_results:
        summary_text += f"A material balance error of {analysis_results['Material Balance Error (%)']:.2f}% was calculated. "
    
    if 'Naphthalene Purity Status' in analysis_results:
        summary_text += analysis_results['Naphthalene Purity Status']
        if 'Naphthalene in Top Product (%)' in analysis_results:
            summary_text += f" (Current purity: {analysis_results['Naphthalene in Top Product (%)']:.2f}%)"
            
    doc.add_paragraph(summary_text)

    # Section 2: Key Performance Indicators
    doc.add_heading('2. Key Performance Indicators (KPIs)', level=1)
    doc.add_paragraph("All values are averages over the analysis period.")
    for key, value in analysis_results.items():
        if key.startswith(('Naphthalene Purity', 'Naphthalene Loss')) or key.endswith('Status') or 'in Top Product' in key or 'in Bottom Product' in key:
            continue
        tag_match = re.search(r'\((.*?)\)', key)
        if tag_match:
            tag = tag_match.group(1)
            unit = TAG_UNITS.get(tag, '')
        else:
            unit = TAG_UNITS.get(key.split(' ')[-1].strip(), '')

        if isinstance(value, str):
            doc.add_paragraph(f"• {key}: {value}")
        else:
            doc.add_paragraph(f"• {key}: {value:.2f} {unit}")
    
    # Section 3: Composition Analysis
    doc.add_heading('3. Composition Analysis', level=1)
    
    doc.add_heading('3.1 Naphthalene Purity', level=2)
    doc.add_paragraph(f"• Naphthalene Purity Status: {analysis_results.get('Naphthalene Purity Status', 'N/A')}")
    doc.add_paragraph(f"• Naphthalene in Top Product (FT-04): {analysis_results.get('Naphthalene in Top Product (%)', 'N/A'):.2f}%")
    doc.add_paragraph(f"• Naphthalene in Bottom Product (FT-07): {analysis_results.get('Naphthalene in Bottom Product (%)', 'N/A'):.2f}%")
    doc.add_paragraph(f"• Naphthalene Loss in C-03 Bottoms: {analysis_results.get('Naphthalene Loss in C-03 Bottoms (%)', 'N/A'):.2f}%")

    doc.add_heading('3.2 Impurity Analysis', level=2)
    doc.add_paragraph(f"• Thianaphthalene in Top Product: {analysis_results.get('Thianaphth. in Top Product (%)', 'N/A'):.2f}%")
    doc.add_paragraph(f"• Quinoline in Top Product: {analysis_results.get('Quinoline in Top Product (ppm)', 'N/A'):.2f} ppm")
    doc.add_paragraph(f"• Unknown Impurities in Top Product: {analysis_results.get('Unknown Impurity in Top Product (%)', 'N/A'):.2f}%")
    
    # Section 4: Performance Plots
    doc.add_heading('4. Performance Plots', level=1)

    doc.add_heading('4.1 Temperature Profile', level=2)
    doc.add_paragraph("The temperature profile plot shows the gradient across the column.")
    doc.add_picture(output_temp_plot_path, width=Inches(6))

    doc.add_heading('4.2 Differential Pressure (DP)', level=2)
    doc.add_paragraph("Differential pressure is a key indicator of flooding or fouling.")
    doc.add_picture(output_dp_plot_path, width=Inches(6))

    doc.add_heading('4.3 Daily Trends', level=2)
    doc.add_paragraph("This plot shows the daily average trends of key variables.")
    doc.add_picture(output_trends_plot_path, width=Inches(6))

    doc.save(output_report_path)
    print(f"Analysis report generated successfully at {output_report_path}")

def main():
    """Main execution function."""
    engine = connect_to_database()
    if engine is None:
        return

    scada_data = get_scada_data(engine)
    if scada_data is None:
        return

    analysis_results, scada_data, outliers = perform_analysis(scada_data)
    
    if analysis_results:
        generate_plots(scada_data)
        generate_word_report(analysis_results, scada_data, outliers)
        print("C-03 analysis complete.")
    else:
        print("Analysis failed: no data to process.")

if __name__ == "__main__":
    main()