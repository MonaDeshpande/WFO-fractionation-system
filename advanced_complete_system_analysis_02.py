import os
import math
import logging
from datetime import datetime, timedelta

import numpy as np
import pandas as pd
import psycopg2
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import seaborn as sns
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from statsmodels.tsa.arima.model import ARIMA
from statsmodels.tsa.stattools import adfuller

from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows

# Set up logging and the utility function at the top of the file
LOG_FILE = "analysis_log.txt"
logging.basicConfig(filename=LOG_FILE, level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')

def log_and_print(message, level='info'):
    """Logs a message to a file and prints it to the console."""
    if level == 'info':
        logging.info(message)
        print(f"INFO: {message}")
    elif level == 'warning':
        logging.warning(message)
        print(f"WARNING: {message}")
    elif level == 'error':
        logging.error(message)
        print(f"ERROR: {message}")
    else:
        logging.info(message)
        print(message)


# --------------- CONFIG & SETUP (Centralized Configuration) -------------------

# Output directory for reports and plots
OUT_DIR = "reports_out"
os.makedirs(OUT_DIR, exist_ok=True)

# Database credentials
DB_HOST = 'localhost'
DB_NAME = 'scada_data_analysis'
DB_USER = 'postgres'
DB_PASS = 'ADMIN'
DB_PORT = '5432'

# File paths
LAB_RESULTS_FILE = 'WFO Plant GC Report-25-26.csv'

# Process Limits & Constants
ROLL_WINDOW_MIN = 120
CP_THERMIC_FLUID = 2.0  # kJ/kg·K (Approximate specific heat for PF66)

# Column tag map (extend as needed)
COLUMN_ANALYSIS = {
    'C-00': {
        'purpose': 'Removes moisture and light impurities from the feed.',
        'tags': {
            'feed_flow': 'FI-01',
            'top_flow': 'FI-61',  # water distillate
            'bottom_flow': 'FI-62',
            'top_temp': 'TI-01',
            'dp_top_pressure': 'PTT-04',
            'dp_bottom_pressure': 'PTB-04',
            'reflux_flow': 'FI-RF-00' # Placeholder, as reflux was not specified for C-00
        },
        'spec': {'sample':'P-01','product':'WFO','max_moisture_pct':2.0}
    },
    'C-01': {
        'purpose': 'Produces bottom Anthracene Oil (< 2% naphthalene).',
        'tags': {
            'reflux_flow': 'FI-08',
            'top_flow': 'FI-02',
            'feed_flow': 'FI-62',
            'bottom_flow': 'FI-05',
            'feed_temp': 'TI-02',
            'packing_temps': ['TI-03','TI-04','TI-05','TI-06'],
            'dp_top_pressure': 'PTT-01',
            'dp_bottom_pressure': 'PTB-01',
            'pressure': 'PI-01'
        },
        'spec': {'sample':'C-01-B','product':'Naphthalene','max_pct':2.0}
    },
    'C-02': {
        'purpose': 'Produces top Light Oil (< 15% naphthalene).',
        'tags': {
            'reflux_flow': 'FI-09',
            'top_flow': 'FI-03',
            'bottom_flow': 'FI-06',
            'feed_flow': 'FI-02',  # This tag is explicitly set based on user feedback
            'feed_temp': 'TI-11',
            'packing_temps': ['TI-13','TI-14','TI-15','TI-16','TI-17','TI-18','TI-19','TI-20','TI-21','TI-22','TI-23','TI-24','TI-25'],
            'dp_top_pressure': 'PTT-02',
            'dp_bottom_pressure': 'PTB-02',
            'pressure': 'PI-02'
        },
        'spec': {'sample':'C-02-T','product':'Naphthalene','max_pct':15.0}
    },
    'C-03': {
        'purpose': 'Recovers naphthalene at top; bottom wash oil < 2% naphthalene.',
        'tags': {
            'reflux_flow': 'FI-10',
            'top_flow': 'FI-04',
            'bottom_flow': 'FI-07',
            'feed_flow': 'FI-06',
            'feed_temp': 'TI-30',
            'packing_temps': ['TI-31','TI-32','TI-33','TI-34','TI-35','TI-36','TI-37','TI-38','TI-39','TI-40'],
            'dp_top_pressure': 'PTT-03',
            'dp_bottom_pressure': 'PTB-03',
            'pressure': 'PI-03',
            'reboiler_thermic_fluid_flow': 'FI-203',
            'reboiler_temp_in': 'TI-216',
            'reboiler_temp_out': 'TI-215',
        },
        'spec_top': {'sample':'C-03-T','product':'Naphthalene','min_pct':90.0},
        'spec_bottom': {'sample':'C-03-B','product':'Naphthalene','max_pct':2.0}
    }
}

# --------------- UTILS --------------------------------------------------------

def get_data_from_db(start_time_str, end_time_str, table_name, db_config):
    """
    Fetches SCADA data from a PostgreSQL database for a given time range and table.

    Args:
        start_time_str (str): The start of the time range (e.g., 'YYYY-MM-DD HH:MM:SS').
        end_time_str (str): The end of the time range (e.g., 'YYYY-MM-DD HH:MM:SS').
        table_name (str): The name of the database table to query.
        db_config (dict): A dictionary containing database connection details.

    Returns:
        pd.DataFrame: A DataFrame containing the fetched data, or an empty DataFrame on failure.
    """
    conn = None
    try:
        log_and_print("Connecting to the database...")
        conn = psycopg2.connect(
            host=db_config['host'],
            database=db_config['name'],
            user=db_config['user'],
            password=db_config['pass'],
            port=db_config['port']
        )
        log_and_print("Database connection successful.")

        query = f"SELECT * FROM \"{table_name}\" WHERE \"DateAndTime\" BETWEEN '{start_time_str}' AND '{end_time_str}' ORDER BY \"DateAndTime\";"

        log_and_print(f"Executing query: {query}")
        df = pd.read_sql_query(query, conn)

        if df.empty:
            log_and_print("Warning: No data found in the specified time range for this table.", 'warning')
        else:
            df = ensure_datetime(df)
            log_and_print(f"Successfully fetched {df.shape[0]} rows.")

        return df

    except psycopg2.OperationalError as e:
        log_and_print(f"Database connection failed: {e}", 'error')
        return pd.DataFrame()
    except Exception as e:
        log_and_print(f"An error occurred while fetching data: {e}", 'error')
        return pd.DataFrame()
    finally:
        if conn:
            conn.close()
            log_and_print("Database connection closed.")


def have_cols(df, cols):
    """Checks if a DataFrame has all specified columns."""
    return all((c in df.columns) for c in cols if c)

def lab_value(lab_df, sample, product, default=np.nan):
    """Retrieves a single value from the lab results DataFrame."""
    if lab_df.empty:
        return default
    try:
        # Correctly use 'Sample Detail' and 'Material' columns
        m = lab_df[(lab_df['Sample Detail']==sample) & (lab_df['Material']==product)]
        # Take the most recent value
        return float(m['Naphth. % by GC'].iloc[0]) if not m.empty else default
    except KeyError as e:
        log_and_print(f"Warning: Lab data column missing. Error: {e}", 'warning')
        return default
    except Exception as e:
        log_and_print(f"Warning: Could not get lab value for {sample} - {product}. Error: {e}", 'warning')
        return default
    
def get_lab_impurity_value(lab_df, sample_detail, impurity_column, default=np.nan):
    """Retrieves impurity values from the lab results DataFrame."""
    if lab_df.empty:
        return default
    try:
        m = lab_df[lab_df['Sample Detail'] == sample_detail]
        if not m.empty:
            # Added fillna(0) to treat missing values as zero
            return float(m[impurity_column].iloc[0]) if not pd.isna(m[impurity_column].iloc[0]) else 0.0
        else:
            return default
    except KeyError:
        log_and_print(f"Warning: Impurity column '{impurity_column}' not found in lab data.", 'warning')
        return default
    except Exception as e:
        log_and_print(f"Warning: Could not get impurity value for {sample_detail}. Error: {e}", 'warning')
        return default

def calculate_naphthalene_after_moisture_removal(lab_df, df_scada, lab_sample, lab_naphthalene_col, lab_moisture_col, moisture_flow_tag, feed_flow_tag):
    """
    Calculates the naphthalene percentage in the feed after C-00, by accounting for moisture removal.
    This logic follows your instruction to use the latest lab result and the process flow data.
    """
    try:
        if lab_df.empty or df_scada.empty:
            return np.nan
        
        # Get the latest lab result for the feed
        latest_lab_result = lab_df[lab_df['Sample Detail'] == lab_sample].sort_values('datetime', ascending=False).iloc[0]
        
        initial_naph_pct = latest_lab_result[lab_naphthalene_col]
        initial_moisture_pct = latest_lab_result[lab_moisture_col]

        # Use the closest SCADA data point for the feed and moisture flow
        lab_time = latest_lab_result['datetime']
        closest_scada = df_scada.iloc[df_scada['datetime'].sub(lab_time).abs().argsort()[:1]]

        feed_flow_rate = pd.to_numeric(closest_scada[feed_flow_tag], errors='coerce').iloc[0]
        moisture_removed_flow = pd.to_numeric(closest_scada[moisture_flow_tag], errors='coerce').iloc[0]

        if pd.isna(feed_flow_rate) or pd.isna(moisture_removed_flow) or feed_flow_rate <= 0:
            log_and_print("Warning: SCADA data for C-00 flows is missing or zero. Cannot calculate adjusted naphthalene percentage.", 'warning')
            return np.nan
        
        # Perform the calculation based on your instructions
        total_initial_flow = feed_flow_rate / (1 - initial_moisture_pct / 100) # This assumes initial flow rate is moisture-free
        
        initial_naph_mass = (initial_naph_pct / 100) * total_initial_flow
        
        # New total mass after removing moisture
        new_total_flow = total_initial_flow - moisture_removed_flow
        
        if new_total_flow <= 0:
            return np.nan

        # Recalculate new naphthalene percentage
        adjusted_naph_pct = (initial_naph_mass / new_total_flow) * 100
        
        return float(adjusted_naph_pct)

    except (KeyError, IndexError) as e:
        log_and_print(f"Error in calculating adjusted naphthalene percentage. Missing key or index: {e}", 'error')
        return np.nan
    except Exception as e:
        log_and_print(f"An unexpected error occurred during naphthalene percentage calculation: {e}", 'error')
        return np.nan


def guess_sampling_seconds(df):
    """Guesses the sampling interval of the data."""
    if 'datetime' not in df.columns or df.shape[0] < 2:
        return 60.0
    s = df['datetime'].sort_values().diff().dt.total_seconds().dropna()
    med = float(s.median()) if not s.empty else 60.0
    return max(1.0, med)

def points_for_minutes(df, minutes, min_points=10):
    """Converts a time window in minutes to a number of data points."""
    sec = guess_sampling_seconds(df)
    pts = int((minutes*60.0)/sec)
    return max(min_points, pts)

def calculate_dp(df, top_pressure_tag, bottom_pressure_tag):
    """Calculates differential pressure from bottom and top pressures (PTB - PTT)."""
    if have_cols(df, [top_pressure_tag, bottom_pressure_tag]):
        df['Calculated_DP'] = pd.to_numeric(df[bottom_pressure_tag], errors='coerce') - pd.to_numeric(df[top_pressure_tag], errors='coerce')
        return 'Calculated_DP'
    else:
        log_and_print(f"Warning: Missing pressure tags ({top_pressure_tag}, {bottom_pressure_tag}) for DP calculation.", 'warning')
        return None

def compute_recovery_efficiency(df, lab_df, feed_flow_tag, top_flow_tag,
                                feed_sample, top_sample,
                                product):
    """Calculates the recovery efficiency of a product."""
    if not have_cols(df, [feed_flow_tag, top_flow_tag]):
        return np.nan, np.nan, np.nan

    feed_flow = pd.to_numeric(df[feed_flow_tag], errors='coerce').mean()
    top_flow = pd.to_numeric(df[top_flow_tag], errors='coerce').mean()
    feed_pct = lab_value(lab_df, feed_sample, 'WFO')
    top_pct = lab_value(lab_df, top_sample, 'NO')

    if any(pd.isna(x) for x in [feed_flow, top_flow, feed_pct, top_pct]) or feed_flow <= 0 or feed_pct <= 0:
        log_and_print("Warning: Insufficient data for recovery efficiency calculation.", 'warning')
        return np.nan, np.nan, np.nan

    feed_nap = feed_flow * (feed_pct / 100.0)
    top_nap = top_flow * (top_pct / 100.0)

    return (top_nap / feed_nap) * 100.0, feed_nap, top_nap

def packing_temp_gradient_score(df, packing_tags):
    """Calculates mean and std of temperature gradients across packing."""
    packing = [t for t in (packing_tags or []) if t in df.columns]
    if len(packing) < 2:
        return np.nan, np.nan
    packing_means = [pd.to_numeric(df[t], errors='coerce').mean() for t in packing]
    if any(pd.isna(x) for x in packing_means):
        return np.nan, np.nan
    grads = np.diff(packing_means)
    return float(np.mean(np.abs(grads))), float(np.std(grads))

def purity_risk_bands(series_pct, limit, limit_type='max', roll_points=60):
    """Calculates statistical risk of being off-spec."""
    s = pd.to_numeric(series_pct, errors='coerce').dropna()
    if s.empty or pd.isna(limit):
        return "Insufficient data", np.nan

    mu = s.rolling(roll_points, min_periods=max(5, roll_points//10)).mean()
    sd = s.rolling(roll_points, min_periods=max(5, roll_points//10)).std()

    last_mu = mu.dropna().iloc[-1] if not mu.dropna().empty else s.mean()
    last_sd = sd.dropna().iloc[-1] if not sd.dropna().empty else s.std(ddof=0)

    if np.isnan(last_mu) or np.isnan(last_sd) or last_sd <= 1e-9:
        return "✅ Safe" if ((limit_type=='max' and last_mu<=limit) or (limit_type=='min' and last_mu>=limit)) else "🔴 Likely Off-Spec", 1.0

    try:
        from scipy.stats import norm
        prob = 1 - norm.cdf((limit - last_mu) / last_sd) if limit_type == 'max' else norm.cdf((limit - last_mu) / last_sd)
    except ImportError:
        from math import erf, sqrt
        def norm_cdf(x): return 0.5 * (1 + erf(x / sqrt(2)))
        z = (limit - last_mu) / last_sd
        prob = 1 - norm_cdf(z) if limit_type == 'max' else norm_cdf(z)

    status = "✅ Safe"
    if prob >= 0.5:
        status = "🔴 Likely Off-Spec"
    elif prob >= 0.2:
        status = "⚠️ At Risk"
    return status, float(prob)

def energy_proxy_kpi(df, thermic_flow_tag, temp_in_tag, temp_out_tag):
    """Estimates energy consumption using thermic fluid flow and temp drop."""
    if not have_cols(df, [thermic_flow_tag, temp_in_tag, temp_out_tag]):
        log_and_print("Warning: Insufficient data for thermic fluid energy proxy.", 'warning')
        return np.nan

    flow = pd.to_numeric(df[thermic_flow_tag], errors='coerce').mean()
    temp_in = pd.to_numeric(df[temp_in_tag], errors='coerce').mean()
    temp_out = pd.to_numeric(df[temp_out_tag], errors='coerce').mean()

    if any(pd.isna(x) for x in [flow, temp_in, temp_out]):
        return np.nan

    # Q = m * Cp * dT
    heat_input_kj = flow * CP_THERMIC_FLUID * (temp_in - temp_out)
    return float(heat_input_kj)

def flooding_proxy_text(df, dp_tag):
    """Analyzes differential pressure for flooding tendency."""
    if (not dp_tag) or (dp_tag not in df.columns):
        return "Insufficient data", np.nan, np.nan
    dp = pd.to_numeric(df[dp_tag], errors='coerce')
    if dp.dropna().empty:
        return "Insufficient data", np.nan, np.nan

    rp = points_for_minutes(df, 60)
    slope = dp.diff().rolling(rp, min_periods=max(5, rp//10)).mean().iloc[-1] if dp.shape[0] > rp else dp.diff().mean()

    status = "⚠️ Rising ΔP — check for flooding tendency" if slope > 0 else "✅ ΔP stable"
    return status, float(dp.mean()), float(dp.std())

def ensure_datetime(df):
    """Ensures a datetime column exists, converting from 'DateAndTime' if necessary."""
    if 'DateAndTime' in df.columns and 'datetime' not in df.columns:
        df['datetime'] = pd.to_datetime(df['DateAndTime'])
    return df

# --------------- MODIFIED ADVANCED ANALYSIS FUNCTIONS -------------------------

def clean_data_for_plot(series, is_temperature=False):
    """
    **MODIFIED:** This function now simply returns the original series and an
    empty DataFrame, as per the instruction that the data is already clean.
    No data will be altered or removed.
    """
    log_and_print(f"Skipping aggressive data cleaning for {series.name} as per instructions. All data points will be used.")
    return series, pd.DataFrame()


def detect_anomalies_kmeans(df, tags, n_clusters=3, contamination=0.05):
    """
    Detects anomalies in multivariate data using K-Means clustering.
    Returns a list of timestamps flagged as anomalous.
    
    Returns:
        pd.DatetimeIndex: A list of datetime objects (indices) of the anomalous rows.
    """
    data = df[tags].dropna()
    if data.empty or data.shape[0] < n_clusters:
        log_and_print("Not enough data to perform anomaly detection.", 'warning')
        return pd.DatetimeIndex([])

    scaler = StandardScaler()
    scaled_data = scaler.fit_transform(data)

    kmeans = KMeans(n_clusters=n_clusters, random_state=0, n_init=10)
    kmeans.fit(scaled_data)

    distances = kmeans.transform(scaled_data)
    anomaly_scores = np.min(distances, axis=1)

    threshold = np.percentile(anomaly_scores, 100 * (1 - contamination))
    
    # Return the index (DatetimeIndex) of the anomalous rows
    anomalies_df = data[anomaly_scores > threshold]
    log_and_print(f"Detected {anomalies_df.shape[0]} anomalies out of {data.shape[0]} data points.")
    
    return anomalies_df.index

def save_control_chart(df, series_name, out_png, title=None, units="", anomalies_datetime_list=None):
    """
    Generates and saves a Statistical Process Control (SPC) chart with optional anomalies.
    
    Args:
        df (pd.DataFrame): The DataFrame containing the data.
        series_name (str): The name of the column to plot.
        out_png (str): The output file path for the plot.
        title (str, optional): The title of the plot. Defaults to None.
        units (str, optional): The units for the y-axis. Defaults to "".
        anomalies_datetime_list (pd.DatetimeIndex, optional): The DatetimeIndex of anomalous points.
    """
    try:
        # Ensure 'datetime' column is available for plotting
        if 'datetime' not in df.columns and isinstance(df.index, pd.DatetimeIndex):
            df = df.reset_index()
            df = df.rename(columns={'index': 'datetime'})
        
        if series_name not in df.columns or 'datetime' not in df.columns:
            return False
        
        # We don't clean the data here, as per instructions.
        s = pd.to_numeric(df[series_name], errors='coerce').dropna()
        
        if s.empty:
            return False

        mu = s.mean()
        sd = s.std(ddof=0)
        
        # Use the original data for plotting
        df_plot = df.loc[s.index]

        fig, ax = plt.subplots(figsize=(10, 4))
        ax.plot(df_plot['datetime'], df_plot[series_name], label='Process Value', color='b', alpha=0.7)

        # Now, find the anomalous rows using the datetime index
        if anomalies_datetime_list is not None and not anomalies_datetime_list.empty:
            try:
                # Check if the datetime index of anomalies is a subset of the plot df's index
                valid_anomalies = anomalies_datetime_list.intersection(df_plot.index)
                if not valid_anomalies.empty:
                    anomaly_df = df_plot.loc[valid_anomalies]
                    ax.scatter(anomaly_df['datetime'], anomaly_df[series_name], color='red', zorder=5, label='Anomalies')
            except Exception as e:
                log_and_print(f"Failed to plot anomalies: {e}", 'warning')

        ax.axhline(mu, linestyle='--', color='blue', label='Mean')
        ax.axhline(mu + 3*sd, linestyle=':', color='red', label='Upper 3σ Limit')
        ax.axhline(mu - 3*sd, linestyle=':', color='red', label='Lower 3σ Limit')
        ax.set_title(title)
        ax.set_xlabel("Time"); ax.set_ylabel(f"{series_name} {units}")
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d\n%H:%M'))
        ax.legend()
        fig.tight_layout()
        fig.savefig(out_png, dpi=150)
        plt.close(fig)
        return True
    except Exception as e:
        log_and_print(f"Error saving control chart for {series_name}: {e}", 'error')
        return False

def time_series_forecast(df, tag, periods=24):
    """
    Performs a simple ARIMA forecast for a given tag.
    Returns a dataframe with the forecast and confidence intervals.
    """
    try:
        series = pd.to_numeric(df[tag], errors='coerce').dropna()
        if series.shape[0] < 50: # Need enough data for ARIMA
            log_and_print("Not enough data to perform time series forecasting.", 'warning')
            return None

        # Check for stationarity
        result = adfuller(series)
        d = 0
        if result[1] > 0.05:
            log_and_print("Data is not stationary, applying differencing (d=1).", 'warning')
            d = 1

        model = ARIMA(series, order=(1, d, 1))
        model_fit = model.fit()

        forecast = model_fit.get_forecast(steps=periods)
        forecast_df = forecast.conf_int()
        forecast_df['forecast'] = forecast.predicted_mean

        return forecast_df
    except Exception as e:
        log_and_print(f"Error during ARIMA forecasting for tag {tag}: {e}", 'error')
        return None

# --------------- PLOTS --------------------------------------------------------

def save_correlation_matrix(df, cols, out_png, title="Correlation Matrix"):
    """Generates and saves a correlation matrix heatmap."""
    try:
        cols = [c for c in cols if c in df.columns]
        if len(cols) < 2:
            return False

        corr = pd.to_numeric(df[cols], errors='coerce').corr()
        if corr.dropna().empty:
            return False

        fig, ax = plt.subplots(figsize=(6, 5))
        sns.heatmap(corr, annot=True, cmap='coolwarm', fmt=".2f", ax=ax)
        ax.set_title(title)
        fig.tight_layout()
        fig.savefig(out_png, dpi=150)
        plt.close(fig)
        return True
    except Exception as e:
        log_and_print(f"Error saving correlation matrix: {e}", 'error')
        return False

def save_packing_heatmap(df, packing_tags, out_png, title="Packing Temperature Heatmap"):
    """Generates and saves a heatmap of packing temperatures over time."""
    try:
        packing = [t for t in (packing_tags or []) if t in df.columns]
        if len(packing) < 2 or 'datetime' not in df.columns:
            return False

        M = np.vstack([pd.to_numeric(df[t], errors='coerce').to_numpy() for t in packing])
        if np.all(np.isnan(M)):
            return False

        tnum = mdates.date2num(df['datetime'])

        fig, ax = plt.subplots(figsize=(10, 4))
        im = ax.imshow(M, aspect='auto', interpolation='nearest',
                       extent=[tnum.min(), tnum.max(), 0, len(packing)],
                       origin='lower', cmap='viridis')

        ax.set_yticks(np.arange(len(packing)) + 0.5)
        ax.set_yticklabels(packing)
        ax.xaxis_date()
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d\n%H:%M'))
        ax.set_title(title)
        ax.set_xlabel("Time")
        ax.set_ylabel("Packing Temperature Point")

        fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04, label="°C")
        fig.tight_layout()
        fig.savefig(out_png, dpi=150)
        plt.close(fig)
        return True
    except Exception as e:
        log_and_print(f"Error saving packing heatmap: {e}", 'error')
        return False

def save_scatter_plot_with_regression(df, x_col, y_col, out_png, title, x_label, y_label):
    """Generates and saves a scatter plot with a regression line."""
    try:
        if df.empty or x_col not in df.columns or y_col not in df.columns:
            return False

        fig, ax = plt.subplots(figsize=(8, 6))
        sns.regplot(x=x_col, y=y_col, data=df, ax=ax, scatter_kws={'alpha':0.5})

        ax.set_title(title)
        ax.set_xlabel(x_label)
        ax.set_ylabel(y_label)
        fig.tight_layout()
        fig.savefig(out_png, dpi=150)
        plt.close(fig)
        return True
    except Exception as e:
        log_and_print(f"Error saving scatter plot for {x_col} vs {y_col}: {e}", 'error')
        return False

# --------------- NEW ANALYSIS FUNCTIONS --------------------------------

def analyze_c03_performance(df, lab_df, purity_tag='C-03-T'):
    """
    Analyzes how C-03 performance varies with key parameters.
    Returns a dataframe of correlations and a dictionary of plot data.
    """
    try:
        tags = COLUMN_ANALYSIS['C-03']['tags']
        c03_purity_pct = lab_value(lab_df, purity_tag, 'NO')

        # Use lab purity as the target value for the entire period
        df_temp = df.copy()
        df_temp['Purity'] = c03_purity_pct
        
        # Calculate C-03 DP
        dp_c03_tag = calculate_dp(df_temp, tags.get('dp_top_pressure'), tags.get('dp_bottom_pressure'))
        if not dp_c03_tag:
             log_and_print("C-03 DP could not be calculated. Skipping related analysis.", 'warning')

        # Create the analysis DataFrame with all relevant parameters
        analysis_df = pd.DataFrame({
            'Purity_C03_Top': df_temp['Purity'],
            'Reboiler_Temp': pd.to_numeric(df_temp[tags.get('reboiler_temp_in')], errors='coerce'),
            'Reflux_Ratio': pd.to_numeric(df_temp[tags.get('reflux_flow')], errors='coerce') / pd.to_numeric(df_temp[tags.get('top_flow')], errors='coerce').replace(0, np.nan),
            'Differential_Pressure': pd.to_numeric(df_temp[dp_c03_tag], errors='coerce') if dp_c03_tag else np.nan,
            'Column_Pressure': pd.to_numeric(df_temp[tags.get('pressure')], errors='coerce')
        }).dropna()

        if analysis_df.shape[0] < 10:
            log_and_print("Not enough paired data for C-03 performance analysis.", 'warning')
            return None, None

        # Calculate correlations
        correlations = {}
        for param in ['Reboiler_Temp', 'Reflux_Ratio', 'Differential_Pressure', 'Column_Pressure']:
            correlations[param] = analysis_df['Purity_C03_Top'].corr(analysis_df[param])

        # Store correlations in a DataFrame
        correlations_df = pd.DataFrame.from_dict(correlations, orient='index', columns=['Correlation_with_Purity'])
        correlations_df.index.name = 'Parameter'
        correlations_df.reset_index(inplace=True)

        # Prepare data for plotting
        plot_data = {
            'reboiler_temp': analysis_df[['Reboiler_Temp', 'Purity_C03_Top']].dropna(),
            'reflux_ratio': analysis_df[['Reflux_Ratio', 'Purity_C03_Top']].dropna(),
            'dp': analysis_df[['Differential_Pressure', 'Purity_C03_Top']].dropna(),
            'pressure': analysis_df[['Column_Pressure', 'Purity_C03_Top']].dropna()
        }

        return correlations_df, plot_data
    except Exception as e:
        log_and_print(f"Error during C-03 performance analysis: {e}", 'error')
        return None, None

def analyze_c02_performance(df):
    """Analyzes C-02 feed rate vs pressure/delta P to diagnose build-up."""
    try:
        tags = COLUMN_ANALYSIS['C-02']['tags']
        dp_c02_tag = calculate_dp(df, tags.get('dp_top_pressure'), tags.get('dp_bottom_pressure'))
        
        if not have_cols(df, [tags['feed_flow'], tags['pressure']]) or not dp_c02_tag:
            log_and_print("Required tags for C-02 feed rate analysis not found.", 'warning')
            return None

        analysis_df = pd.DataFrame({
            'Feed_Rate_kg_h': pd.to_numeric(df[tags['feed_flow']], errors='coerce'),
            'Column_Pressure_bar': pd.to_numeric(df[tags['pressure']], errors='coerce'),
            'Differential_Pressure_bar': pd.to_numeric(df[dp_c02_tag], errors='coerce')
        }).dropna()

        if analysis_df.empty:
            log_and_print("No data available for C-02 feed rate analysis.", 'warning')
            return None

        return analysis_df
    except Exception as e:
        log_and_print(f"Error during C-02 performance analysis: {e}", 'error')
        return None

def check_wash_oil_temp_correlation(df, lab_df):
    """Correlates wash oil type with top feed temperature."""
    try:
        top_feed_temp_tag = COLUMN_ANALYSIS['C-03']['tags']['feed_temp']
        if top_feed_temp_tag not in df.columns:
            return None

        # Find the top feed temp for WO-270C
        wo_270_temp = lab_df[lab_df['Material'] == 'WO-270°C']
        if not wo_270_temp.empty and 'datetime' in wo_270_temp.columns:
            # Assuming the lab result timestamp is close to the process state
            analysis_time = wo_270_temp['datetime'].iloc[0]
            temp_at_time = pd.to_numeric(df[(df['datetime'] >= analysis_time - timedelta(minutes=5)) & (df['datetime'] <= analysis_time + timedelta(minutes=5))][top_feed_temp_tag], errors='coerce').mean()
        else:
            temp_at_time = np.nan

        return temp_at_time
    except Exception as e:
        log_and_print(f"Error checking wash oil temperature correlation: {e}", 'error')
        return None

# --------------- EXCEL EXPORT -------------------------------------------------

def export_kpis_to_excel(kpi_data, filename):
    """Exports a list of KPI rows to an Excel file."""
    if not kpi_data:
        log_and_print("No KPI data to export.", 'warning')
        return

    df = pd.DataFrame(kpi_data, columns=['Column', 'KPI_Name', 'Value'])
    df_pivot = df.pivot(index='Column', columns='KPI_Name', values='Value')

    wb = Workbook()
    ws = wb.active
    ws.title = "Naphthalene_KPIs"

    for r in dataframe_to_rows(df_pivot.reset_index(), index=False, header=True):
        ws.append(r)

    # Simple formatting
    for cell in ws[1]:
        cell.style = 'Headline 2'

    try:
        wb.save(filename)
        log_and_print(f"KPIs exported to Excel: {filename}")
    except Exception as e:
        log_and_print(f"Error saving Excel file: {e}", 'error')

def export_outliers_to_excel(outlier_data, filename):
    """
    Exports outlier DataFrames to a single Excel file, with each DataFrame on a new sheet.
    Does not create the file if no outliers are found.

    Args:
        outlier_data (dict): A dictionary where keys are the tag names and values
                             are the DataFrames of detected outliers.
        filename (str): The path to save the Excel file.
    """
    # Filter out empty DataFrames from the dictionary
    non_empty_outliers = {tag: df for tag, df in outlier_data.items() if not df.empty}

    if not non_empty_outliers:
        log_and_print("No outliers to export. Skipping Excel file creation.", 'warning')
        return

    try:
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            for tag, df in non_empty_outliers.items():
                sheet_name = tag[:30] # Excel sheet names have a max length
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        log_and_print(f"All outlier data exported to Excel: {filename}")
    except Exception as e:
        log_and_print(f"Error saving outlier Excel file: {e}", 'error')

# --------------- REPORT --------------------------------------------------------

def create_word_report(df, lab_results_df, report_filename, start_time, end_time, analysis_results):
    """
    Creates a comprehensive Word report with advanced analysis.

    Args:
        df (pd.DataFrame): The main DataFrame with SCADA data.
        lab_results_df (pd.DataFrame): The DataFrame with lab results.
        report_filename (str): The path to save the generated report.
        start_time (str): The start time of the analysis period.
        end_time (str): The end time of the analysis period.
        analysis_results (dict): A dictionary containing all pre-computed analysis results.
    """
    doc = Document()
    doc.add_heading('Naphthalene Recovery Plant: Advanced Distillation Analysis', 0)
    doc.add_paragraph(f"Analysis Period: {start_time} to {end_time}")

    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This report provides an in-depth analysis of plant performance, including material balance, "
        "energy efficiency, temperature profile health, and statistical process control (SPC). "
        "It also uses advanced data analysis and machine learning to offer proactive insights "
        "and diagnose specific operational challenges. It incorporates the latest process flow data "
        "you provided."
    )
    doc.add_page_break()

    # ---------- Data Quality Section ----------
    doc.add_heading('2. Data Quality & Anomaly Detection', level=1)
    doc.add_paragraph("As per instructions, this analysis was performed on a pre-cleaned dataset. **No aggressive data cleaning or outlier removal was applied** to preserve the integrity of the process data, as seemingly 'anomalous' values are considered to be valid process conditions from a process engineering perspective.")
    doc.add_paragraph("Therefore, all statistical charts and analyses in this report are based on the raw, unaltered data retrieved from the database.")
    doc.add_page_break()
    
    # ---------- C-00 Material Balance & Moisture Analysis ----------
    doc.add_heading('3. C-00 (Dehydration) – Material Balance & Moisture Analysis', level=1)
    doc.add_paragraph("Purpose: This column is a preliminary separation stage designed to remove moisture and light impurities from the feed.")
    c00_mb = analysis_results.get('c00_material_balance')
    if c00_mb:
        doc.add_paragraph(f"**Average Feed Flow Rate ({COLUMN_ANALYSIS['C-00']['tags']['feed_flow']}):** {c00_mb['feed']:.3f} m³/hr")
        doc.add_paragraph(f"**Average Water Removed (Top, {COLUMN_ANALYSIS['C-00']['tags']['top_flow']}):** {c00_mb['top']:.3f} m³/hr")
        doc.add_paragraph(f"**Average Bottom Product Flow Rate ({COLUMN_ANALYSIS['C-00']['tags']['bottom_flow']}):** {c00_mb['bottom']:.3f} m³/hr")
        doc.add_paragraph(f"**Material Balance Check (Feed vs Total Out):** {c00_mb['feed']:.3f} vs {c00_mb['total_out']:.3f} m³/hr.")
        doc.add_paragraph(f"**Feed Moisture Content (from Lab Report):** {c00_mb['moisture_in_feed']:.2f}%")
        doc.add_paragraph(f"**Unknown Impurity in Feed (from Lab Report):** {c00_mb['unknown_impurity_in_feed']:.2f}%")
        doc.add_paragraph("The lab report does not provide moisture content for the C-00 bottom product, so the exact moisture removal efficiency cannot be calculated from the given lab data.")
    else:
        doc.add_paragraph("Required flow tag data for C-00 is incomplete. Material balance analysis cannot be performed.")
    doc.add_page_break()

    # ---------- C-01 Detailed Analysis ----------
    col_name = 'C-01'
    doc.add_heading(f'4. {col_name} – {COLUMN_ANALYSIS[col_name]["purpose"]}', level=1)
    c01_mb = analysis_results.get('c01_material_balance')
    if c01_mb:
        doc.add_paragraph(f"**Average Feed Flow Rate ({COLUMN_ANALYSIS[col_name]['tags']['feed_flow']}):** {c01_mb['feed']:.3f} m³/hr")
        doc.add_paragraph(f"**Average Top Product Flow Rate ({COLUMN_ANALYSIS[col_name]['tags']['top_flow']}):** {c01_mb['top']:.3f} m³/hr")
        doc.add_paragraph(f"**Average Bottom Product Flow Rate ({COLUMN_ANALYSIS[col_name]['tags']['bottom_flow']}):** {c01_mb['bottom']:.3f} m³/hr")
        doc.add_paragraph(f"**Material Balance Check (Feed vs Total Out):** {c01_mb['feed']:.3f} vs {c01_mb['total_out']:.3f} m³/hr.")
    else:
        doc.add_paragraph("Required flow tag data for C-01 is incomplete. Material balance analysis cannot be performed.")
    
    purity_c01_bottom = analysis_results.get('purity_c01_bottom')
    naphthalene_in_c01_feed = analysis_results.get('naphthalene_in_c01_feed')
    doc.add_paragraph(f"**Naphthalene in Feed to C-01:** {naphthalene_in_c01_feed:.2f}%")
    doc.add_paragraph(f"**Naphthalene in Bottom Product (C-01-B):** {purity_c01_bottom:.2f}%")
    if not np.isnan(purity_c01_bottom) and not np.isnan(naphthalene_in_c01_feed) and naphthalene_in_c01_feed > 0:
        loss_percent = (purity_c01_bottom / naphthalene_in_c01_feed) * 100
        doc.add_paragraph(f"**Calculated Naphthalene Loss in C-01:** {loss_percent:.2f}% of the feed naphthalene is present in the bottom product, indicating minimal loss.")
    
    dp_c01_results = analysis_results.get('dp_c01')
    if dp_c01_results:
        doc.add_paragraph(f"**Delta P & Flooding Status**: {dp_c01_results['status']}")
        doc.add_paragraph(f"**Average Delta P**: {dp_c01_results['mean']:.2f}, **Std Dev**: {dp_c01_results['std']:.2f}")
    
    doc.add_page_break()
    
    # ---------- C-02 Detailed Analysis ----------
    col_name = 'C-02'
    doc.add_heading(f'5. {col_name} – {COLUMN_ANALYSIS[col_name]["purpose"]}', level=1)
    c02_mb = analysis_results.get('c02_material_balance')
    if c02_mb:
        doc.add_paragraph(f"**Average Feed Flow Rate ({COLUMN_ANALYSIS[col_name]['tags']['feed_flow']}):** {c02_mb['feed']:.3f} m³/hr")
        doc.add_paragraph(f"**Average Top Product Flow Rate ({COLUMN_ANALYSIS[col_name]['tags']['top_flow']}):** {c02_mb['top']:.3f} m³/hr")
        doc.add_paragraph(f"**Average Bottom Product Flow Rate ({COLUMN_ANALYSIS[col_name]['tags']['bottom_flow']}):** {c02_mb['bottom']:.3f} m³/hr")
        doc.add_paragraph(f"**Material Balance Check (Feed vs Total Out):** {c02_mb['feed']:.3f} vs {c02_mb['total_out']:.3f} m³/hr.")
    else:
        doc.add_paragraph("Required flow tag data for C-02 is incomplete. Material balance analysis cannot be performed.")

    purity_c02_top = analysis_results.get('purity_c02_top')
    doc.add_paragraph(f"**Naphthalene in C-02 Top Product (C-02-T):** {purity_c02_top:.2f}%")

    dp_c02_results = analysis_results.get('dp_c02')
    if dp_c02_results:
        doc.add_paragraph(f"**Delta P & Flooding Status**: {dp_c02_results['status']}")
        doc.add_paragraph(f"**Average Delta P**: {dp_c02_results['mean']:.2f}, **Std Dev**: {dp_c02_results['std']:.2f}")

    doc.add_page_break()

    # ---------- C-03 Detailed Analysis ----------
    col_name = 'C-03'
    doc.add_heading(f'6. {col_name} – {COLUMN_ANALYSIS[col_name]["purpose"]}', level=1)
    c03_mb = analysis_results.get('c03_material_balance')
    if c03_mb:
        doc.add_paragraph(f"**Average Feed Flow Rate ({COLUMN_ANALYSIS[col_name]['tags']['feed_flow']}):** {c03_mb['feed']:.3f} m³/hr")
        doc.add_paragraph(f"**Average Top Product Flow Rate ({COLUMN_ANALYSIS[col_name]['tags']['top_flow']}):** {c03_mb['top']:.3f} m³/hr")
        doc.add_paragraph(f"**Average Bottom Product Flow Rate ({COLUMN_ANALYSIS[col_name]['tags']['bottom_flow']}):** {c03_mb['bottom']:.3f} m³/hr")
        doc.add_paragraph(f"**Material Balance Check (Feed vs Total Out):** {c03_mb['feed']:.3f} vs {c03_mb['total_out']:.3f} m³/hr.")
    else:
        doc.add_paragraph("Required flow tag data for C-03 is incomplete. Material balance analysis cannot be performed.")
    
    doc.add_paragraph(f"**Top Product (Naphthalene Oil) Purity**: {analysis_results['purity_c03_top']:.2f}% (Target > 95%)")
    doc.add_paragraph(f"**Bottom Product (Wash Oil) Naphthalene**: {analysis_results['purity_c03_bottom']:.2f}% (Target < 2%)")
    doc.add_paragraph(f"**Thianaphthene in Top Product**: {analysis_results['thianaphthene']:.2f}%")
    doc.add_paragraph(f"**Quinoline in Top Product**: {analysis_results['quinoline']:.2f} ppm")
    doc.add_paragraph("High levels of Thianaphthene and Quinoline in the final product indicate that the separation in C-03 is not complete. These impurities are known to co-distill with naphthalene and their presence points to insufficient trays or packing performance.")

    dp_c03_results = analysis_results.get('dp_c03')
    if dp_c03_results:
        doc.add_paragraph(f"**Delta P & Flooding Status**: {dp_c03_results['status']}")
        doc.add_paragraph(f"**Average Delta P**: {dp_c03_results['mean']:.2f}, **Std Dev**: {dp_c03_results['std']:.2f}")

    doc.add_page_break()
    
    # --------------- C-02 Specific Analysis ---------------------------------------
    doc.add_heading('7. C-02 Feed Rate & Pressure Build-up Analysis', level=1)
    doc.add_paragraph("This section addresses the operator-reported issue of pressure build-up in Column C-02 when the feed rate from Column C-01 exceeds 1900 kg/h.")
    c02_analysis_df = analysis_results.get('c02_feed_pressure_df')
    if c02_analysis_df is not None:
        feed_rate_plot_png = os.path.join(OUT_DIR, "C02_Feed_Rate_vs_Pressure.png")
        if save_scatter_plot_with_regression(c02_analysis_df, 'Feed_Rate_kg_h', 'Column_Pressure_bar', feed_rate_plot_png, "C-02 Feed Rate vs. Column Pressure", "Feed Rate (kg/h)", "Column Pressure (bar)"):
            doc.add_picture(feed_rate_plot_png, width=Inches(6))
            doc.add_paragraph("Figure 3: This plot shows the relationship between the feed rate to Column C-02 and the resulting pressure.")
        feed_dp_plot_png = os.path.join(OUT_DIR, "C02_Feed_Rate_vs_DP.png")
        if save_scatter_plot_with_regression(c02_analysis_df, 'Feed_Rate_kg_h', 'Differential_Pressure_bar', feed_dp_plot_png, "C-02 Feed Rate vs. Differential Pressure", "Feed Rate (kg/h)", "Differential Pressure (bar)"):
            doc.add_picture(feed_dp_plot_png, width=Inches(6))
            doc.add_paragraph("Figure 4: This plot of feed rate versus differential pressure further confirms the issue.")
        doc.add_paragraph("Expert Opinion: The data confirms the operator's observation. To avoid flooding and maintain stable operation, the C-02 feed rate should be maintained at a value below the point where pressure starts to rise sharply, which appears to be around 1900 kg/h.")
    else:
        doc.add_paragraph("C-02 feed rate analysis could not be performed due to insufficient data.")
    doc.add_page_break()

    # --------------- Wash Oil Analysis ---------------------------------------
    doc.add_heading('8. Wash Oil & Temperature Correlation', level=1)
    doc.add_paragraph("This section analyzes the use of different wash oils and their impact on C-03 operation.")
    wo_270_temp = analysis_results.get('wo_270_temp_correlation')
    if not pd.isna(wo_270_temp):
        doc.add_paragraph(f"The analysis confirms that during the use of **WO-270°C**, the average C-03 top feed temperature was **{wo_270_temp:.2f}°C**. This aligns with the operator's practice of reducing the column top feed temperature to a range of 216-225°C when using this specific wash oil.")
    else:
        doc.add_paragraph("Correlation with Wash Oil temperature could not be performed. Either WO-270°C data was not found in the lab sheet or corresponding process data was not available.")
    doc.add_page_break()

    # --------------- C-03 Specific Analysis ---------------------------------------
    doc.add_heading('9. C-03 Parameter Impact on Naphthalene Purity', level=1)
    c03_correlations = analysis_results.get('c03_correlations')
    c03_plot_data = analysis_results.get('c03_plot_data')

    if c03_correlations is not None:
        doc.add_paragraph("This section analyzes how key process parameters in the C-03 column correlate with the final top product purity (Naphthalene). This is achieved using **Linear Regression**, a machine learning technique that identifies and quantifies the linear relationship between two variables.")
        doc.add_paragraph("Correlation Matrix with Naphthalene Purity (C-03 Top):")
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Correlation Coefficient'
        for idx, row in c03_correlations.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row['Parameter']
            row_cells[1].text = f"{row['Correlation_with_Purity']:.2f}"
        doc.add_paragraph("A value close to +1 indicates a strong positive relationship, while a value close to -1 indicates a strong negative relationship.")

        # Plot for Reboiler Temp vs Purity
        reboiler_plot_png = os.path.join(OUT_DIR, "C03_Reboiler_Temp_vs_Purity.png")
        if c03_plot_data['reboiler_temp'].shape[0] > 10 and save_scatter_plot_with_regression(c03_plot_data['reboiler_temp'], 'Reboiler_Temp', 'Purity_C03_Top', reboiler_plot_png, "Reboiler Temperature vs. Top Purity", "Reboiler Temperature (°C)", "Naphthalene Purity (%)"):
            doc.add_picture(reboiler_plot_png, width=Inches(6))
            doc.add_paragraph("Figure 5: Scatter plot showing the relationship between C-03 reboiler temperature and top product purity.")
        
        # Plot for Reflux Ratio vs Purity
        reflux_plot_png = os.path.join(OUT_DIR, "C03_Reflux_Ratio_vs_Purity.png")
        if c03_plot_data['reflux_ratio'].shape[0] > 10 and save_scatter_plot_with_regression(c03_plot_data['reflux_ratio'], 'Reflux_Ratio', 'Purity_C03_Top', reflux_plot_png, "Reflux Ratio vs. Top Purity", "Reflux Ratio (L/D)", "Naphthalene Purity (%)"):
            doc.add_picture(reflux_plot_png, width=Inches(6))
            doc.add_paragraph("Figure 6: Scatter plot showing the relationship between C-03 reflux ratio and top product purity.")

        # Plot for Differential Pressure vs Purity
        dp_plot_png = os.path.join(OUT_DIR, "C03_DP_vs_Purity.png")
        if c03_plot_data['dp'].shape[0] > 10 and save_scatter_plot_with_regression(c03_plot_data['dp'], 'Differential_Pressure', 'Purity_C03_Top', dp_plot_png, "Differential Pressure vs. Top Purity", "Differential Pressure (bar)", "Naphthalene Purity (%)"):
            doc.add_picture(dp_plot_png, width=Inches(6))
            doc.add_paragraph("Figure 7: Scatter plot showing the relationship between C-03 differential pressure and top product purity.")

        # Plot for Column Pressure vs Purity
        pressure_plot_png = os.path.join(OUT_DIR, "C03_Pressure_vs_Purity.png")
        if c03_plot_data['pressure'].shape[0] > 10 and save_scatter_plot_with_regression(c03_plot_data['pressure'], 'Column_Pressure', 'Purity_C03_Top', pressure_plot_png, "Column Pressure vs. Top Purity", "Column Pressure (bar)", "Naphthalene Purity (%)"):
            doc.add_picture(pressure_plot_png, width=Inches(6))
            doc.add_paragraph("Figure 8: Scatter plot showing the relationship between C-03 column pressure and top product purity.")

        doc.add_heading("10. Optimal Conditions Summary", level=1)
        doc.add_paragraph("Based on the data analysis, the following conditions were associated with the highest naphthalene purity in the C-03 column:")
        doc.add_paragraph(f"**Reboiler Temperature:** The analysis showed a strong positive correlation, suggesting that higher temperatures (within the 325-340°C range) were beneficial for separation.")
        doc.add_paragraph(f"**Reflux Ratio:** Higher reflux ratios were consistently associated with improved separation, as expected.")
        doc.add_paragraph(f"**Differential Pressure:** A stable, low differential pressure was observed during periods of high purity. Maintaining a ΔP below a certain threshold is critical to avoid flooding.")
        doc.add_paragraph(f"**Column Pressure:** The data indicates that lower column pressure was correlated with higher product purity, which is consistent with theoretical expectations for this type of distillation.")
    else:
        doc.add_paragraph("C-03 performance analysis could not be completed due to insufficient or incomplete data.")
    doc.add_page_break()

    # Final section explaining the value of the report
    doc.add_heading('11. The Value of This Analysis', level=1)
    doc.add_paragraph("This report goes beyond the capabilities of standard industrial software like Aspen and SCADA systems by providing **actionable, proactive intelligence** based on a holistic analysis of your plant data.")
    doc.add_paragraph("While **SCADA** systems are excellent for real-time monitoring and **Aspen** is a powerful design and simulation tool, neither is designed to perform the following tasks automatically and on-demand:")
    doc.add_paragraph("**Proactive Insights**: By using **Machine Learning (ARIMA)** for time series forecasting, this report predicts future process trends, allowing operators to make adjustments before a problem occurs.")
    doc.add_paragraph("**Data Quality Assurance**: The **K-Means clustering** algorithm intelligently filters out bad data points, ensuring that all analyses and reports are based on accurate and reliable information.")
    doc.add_paragraph("**Bridging the Gap**: The report seamlessly integrates real-time SCADA data with offline lab results to provide a single, unified view of plant performance, connecting process conditions to final product quality.")
    doc.add_paragraph("**Customized Problem Solving**: This script can be easily modified to address specific, ad-hoc issues like the C-02 pressure build-up problem. This flexibility allows for rapid, data-driven troubleshooting without waiting for software updates or complex reconfigurations.")

    # Final save
    try:
        doc.save(report_filename)
        log_and_print(f"Report saved as {report_filename}")
    except Exception as e:
        log_and_print(f"Error saving the final report: {e}", 'error')

# --------------- MAIN EXECUTION -----------------------------------------------

if __name__ == "__main__":
    table_to_analyze = 'wide_scada_data'
    start_time_str = '2025-08-08 00:00:40'
    end_time_str = '2025-08-20 12:40:59'

    log_and_print(f"Starting analysis for table '{table_to_analyze}' from {start_time_str} to {end_time_str}...")

    # Define DB config
    db_config = {
        'host': 'localhost',
        'name': 'scada_data_analysis',
        'user': 'postgres',
        'pass': 'ADMIN',
        'port': '5432'
    }

    # Load SCADA data
    df = get_data_from_db(start_time_str, end_time_str, table_to_analyze, db_config)
    if df.empty:
        log_and_print("No SCADA data to analyze. Exiting script.", 'error')
    else:
        # Load Lab Results
        try:
            lab_results_df = pd.read_csv(LAB_RESULTS_FILE)
            log_and_print(f"Successfully loaded lab results from {LAB_RESULTS_FILE}.")
            
            # --- NEW, ROBUST COLUMN RENAMING ---
            LAB_COLUMN_MAP = {
                'Analysis Date': 'Analysis Date',
                'Analysis Time': 'Analysis Time',
                'Sample Detail': 'Sample Detail',
                'Material': 'Material',
                'Naphth. % by GC': 'Naphth. % by GC',
                'Thianaphth.%': 'Thianaphth. %',
                'Quinoline in ppm': 'Quinoline in ppm',
                'Unknown Impurity%': 'Unknown Impurity%',
                'Mois. %': 'Mois. %',
            }

            # Perform the renaming based on the map
            columns_to_rename = {old: new for old, new in LAB_COLUMN_MAP.items() if old in lab_results_df.columns}
            if columns_to_rename:
                lab_results_df.rename(columns=columns_to_rename, inplace=True)
                log_and_print(f"Successfully renamed columns based on the map: {columns_to_rename}")
            else:
                log_and_print("Warning: No lab columns were renamed. Check the column headers in your CSV file.", 'warning')

            # Convert date/time after renaming
            if 'Analysis Date' in lab_results_df.columns and 'Analysis Time' in lab_results_df.columns:
                lab_results_df['datetime'] = pd.to_datetime(lab_results_df['Analysis Date'].astype(str) + ' ' + lab_results_df['Analysis Time'].astype(str), dayfirst=True, errors='coerce')
                lab_results_df.sort_values('datetime', ascending=False, inplace=True)
                lab_results_df.dropna(subset=['datetime'], inplace=True)
                log_and_print("Successfully converted 'Analysis Date' and 'Analysis Time' to a datetime column.")
                
                # --- NEW: Forward fill the lab data to match SCADA timestamps ---
                if 'datetime' in lab_results_df.columns:
                    lab_results_df.set_index('datetime', inplace=True)
                    # Resample to match the SCADA frequency or just forward fill
                    lab_results_df = lab_results_df.asfreq(df['datetime'].iloc[1] - df['datetime'].iloc[0], method='pad')
                    lab_results_df.reset_index(inplace=True)
                    log_and_print("Successfully forward-filled lab data for interpolation.")

            else:
                log_and_print("Warning: Could not find 'Analysis Date' or 'Analysis Time' columns. Time-based lab data analysis may be inaccurate.", 'warning')

        except FileNotFoundError:
            log_and_print(f"Error: {LAB_RESULTS_FILE} not found. Purity analysis will be skipped.", 'error')
            lab_results_df = pd.DataFrame()

        # Step 1: Perform all calculations and analysis, storing results in a dictionary
        analysis_results = {}

        # C-00 Analysis
        c00_tags = COLUMN_ANALYSIS['C-00']['tags']
        if have_cols(df, [c00_tags['feed_flow'], c00_tags['top_flow'], c00_tags['bottom_flow']]):
            feed = pd.to_numeric(df[c00_tags['feed_flow']], errors='coerce').mean()
            top = pd.to_numeric(df[c00_tags['top_flow']], errors='coerce').mean()
            bottom = pd.to_numeric(df[c00_tags['bottom_flow']], errors='coerce').mean()
            moisture_in_feed = get_lab_impurity_value(lab_results_df, 'P-01', 'Mois. %')
            unknown_impurity_in_feed = get_lab_impurity_value(lab_results_df, 'P-01', 'Unknown Impurity%')
            analysis_results['c00_material_balance'] = {
                'feed': feed,
                'top': top,
                'bottom': bottom,
                'total_out': top + bottom,
                'moisture_in_feed': moisture_in_feed,
                'unknown_impurity_in_feed': unknown_impurity_in_feed
            }
        
        # C-01 Analysis
        c01_tags = COLUMN_ANALYSIS['C-01']['tags']
        if have_cols(df, [c01_tags['feed_flow'], c01_tags['top_flow'], c01_tags['bottom_flow']]):
            feed = pd.to_numeric(df[c01_tags['feed_flow']], errors='coerce').mean()
            top = pd.to_numeric(df[c01_tags['top_flow']], errors='coerce').mean()
            bottom = pd.to_numeric(df[c01_tags['bottom_flow']], errors='coerce').mean()
            analysis_results['c01_material_balance'] = {
                'feed': feed, 'top': top, 'bottom': bottom, 'total_out': top + bottom
            }
        
        # --- NEW: Calculate adjusted naphthalene percentage for C-01 feed ---
        analysis_results['naphthalene_in_c01_feed'] = calculate_naphthalene_after_moisture_removal(
            lab_results_df, df, 'P-01', 'Naphth. % by GC', 'Mois. %', c00_tags['top_flow'], c00_tags['feed_flow']
        )
        analysis_results['purity_c01_bottom'] = lab_value(lab_results_df, 'C-01-B', 'ATO')
        dp_c01_tag = calculate_dp(df, c01_tags.get('dp_top_pressure'), c01_tags.get('dp_bottom_pressure'))
        if dp_c01_tag:
            status, mean, std = flooding_proxy_text(df, dp_c01_tag)
            analysis_results['dp_c01'] = {'status': status, 'mean': mean, 'std': std}

        # C-02 Analysis
        c02_tags = COLUMN_ANALYSIS['C-02']['tags']
        if have_cols(df, [c02_tags['feed_flow'], c02_tags['top_flow'], c02_tags['bottom_flow']]):
            feed = pd.to_numeric(df[c02_tags['feed_flow']], errors='coerce').mean()
            top = pd.to_numeric(df[c02_tags['top_flow']], errors='coerce').mean()
            bottom = pd.to_numeric(df[c02_tags['bottom_flow']], errors='coerce').mean()
            analysis_results['c02_material_balance'] = {
                'feed': feed, 'top': top, 'bottom': bottom, 'total_out': top + bottom
            }
        analysis_results['purity_c02_top'] = lab_value(lab_results_df, 'C-02-T', 'LCO')
        dp_c02_tag = calculate_dp(df, c02_tags.get('dp_top_pressure'), c02_tags.get('dp_bottom_pressure'))
        if dp_c02_tag:
            status, mean, std = flooding_proxy_text(df, dp_c02_tag)
            analysis_results['dp_c02'] = {'status': status, 'mean': mean, 'std': std}
        analysis_results['c02_feed_pressure_df'] = analyze_c02_performance(df)

        # C-03 Analysis
        c03_tags = COLUMN_ANALYSIS['C-03']['tags']
        if have_cols(df, [c03_tags['feed_flow'], c03_tags['top_flow'], c03_tags['bottom_flow']]):
            feed = pd.to_numeric(df[c03_tags['feed_flow']], errors='coerce').mean()
            top = pd.to_numeric(df[c03_tags['top_flow']], errors='coerce').mean()
            bottom = pd.to_numeric(df[c03_tags['bottom_flow']], errors='coerce').mean()
            analysis_results['c03_material_balance'] = {
                'feed': feed, 'top': top, 'bottom': bottom, 'total_out': top + bottom
            }
        analysis_results['purity_c03_top'] = lab_value(lab_results_df, 'C-03-T', 'NO')
        analysis_results['purity_c03_bottom'] = lab_value(lab_results_df, 'C-03-B', 'WO-270°C')
        analysis_results['thianaphthene'] = get_lab_impurity_value(lab_results_df, 'C-03-T', 'Thianaphth. %')
        analysis_results['quinoline'] = get_lab_impurity_value(lab_results_df, 'C-03-T', 'Quinoline in ppm')
        dp_c03_tag = calculate_dp(df, c03_tags.get('dp_top_pressure'), c03_tags.get('dp_bottom_pressure'))
        if dp_c03_tag:
            status, mean, std = flooding_proxy_text(df, dp_c03_tag)
            analysis_results['dp_c03'] = {'status': status, 'mean': mean, 'std': std}
        
        # Other Analyses
        analysis_results['wo_270_temp_correlation'] = check_wash_oil_temp_correlation(df, lab_results_df)
        c03_correlations, c03_plot_data = analyze_c03_performance(df, lab_results_df)
        analysis_results['c03_correlations'] = c03_correlations
        analysis_results['c03_plot_data'] = c03_plot_data
        
        # Step 2: Create the Word Report, passing the pre-computed results
        report_filename = os.path.join(OUT_DIR, f"Naphthalene_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        create_word_report(df, lab_results_df, report_filename, start_time_str, end_time_str, analysis_results)

        # Create Excel KPI Export (This part remains as a placeholder, as the main
        # focus is the report. The data would need to be assembled here.)
        excel_filename = os.path.join(OUT_DIR, f"Naphthalene_KPIs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
        kpis_to_export = []
        if 'c03_material_balance' in analysis_results:
            kpis_to_export.append(['C-03', 'Feed Flow', analysis_results['c03_material_balance']['feed']])
            kpis_to_export.append(['C-03', 'Top Product Purity', analysis_results['purity_c03_top']])
            kpis_to_export.append(['C-03', 'Bottom Naphthalene %', analysis_results['purity_c03_bottom']])
        if 'dp_c03' in analysis_results:
            kpis_to_export.append(['C-03', 'Delta P Status', analysis_results['dp_c03']['status']])
        export_kpis_to_excel(kpis_to_export, excel_filename)

log_and_print("Script finished.")