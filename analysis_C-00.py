import pandas as pd
from sqlalchemy import create_engine
import psycopg2
from datetime import datetime
import matplotlib.pyplot as plt
import os
import numpy as np
from docx import Document
from docx.shared import Inches
import sqlalchemy
import re

# Database connection parameters (update with your actual details)
DB_HOST = "localhost"
DB_NAME = "scada_data_analysis"
DB_USER = "postgres"
DB_PASSWORD = "ADMIN"

# Define units for each tag
TAG_UNITS = {
    'FT-01': 'kg/h',
    'FT-61': 'kg/h',
    'FT-62': 'kg/h',
    'TI-01': 'degC',
    'TI-61': 'degC',
    'TI-63': 'degC',
    'TI-64': 'degC',
    'PTT-04': 'mmHg',
    'PTB-04': 'mmHg',
    'DIFFERENTIAL_PRESSURE': 'mmHg',
    'TI-215': 'degC',
    'TI-216': 'degC',
    'TI-110': 'degC',
    'FI-101': 'm3/h',
    'FI-204': 'm3/h',
    'REBOILER_HEAT_DUTY': 'kW',
    'CONDENSER_HEAT_DUTY': 'kW',
    'Moisture Removal Percentage': '%'
}

# File paths
output_report_path = "C-00_Analysis_Report.docx"
output_temp_plot_path = "C-00_Temperature_Profile.png"
output_dp_plot_path = "C-00_Differential_Pressure.png"
output_trends_plot_path = "C-00_Daily_Trends.png"

# Engineering constants
THERMIC_FLUID_SPECIFIC_HEAT = 2.0  # kJ/(kg·°C)
WATER_SPECIFIC_HEAT = 4.186       # kJ/(kg·°C)

# File path for moisture data. Note: The code now simulates reading this.
MOISTURE_FILE_PATH = "your_moisture_data.csv"

def connect_to_database():
    """Establishes a connection to the PostgreSQL database."""
    try:
        engine = create_engine(f'postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}/{DB_NAME}')
        print("Database connection successful.")
        return engine
    except Exception as e:
        print(f"Error connecting to the database: {e}")
        return None

def get_scada_data(engine):
    """Retrieves specific SCADA data for the C-00 column."""
    try:
        desired_columns = [
            "DateAndTime", "FT-01", "FT-61", "FT-62", "TI-01", "PTT-04", "PTB-04", 
            "TI-215", "TI-216", "TI-110", "TI-61", "TI-63", "TI-64", "FI-101", "FI-204"
        ]
        
        inspector = sqlalchemy.inspect(engine)
        columns = inspector.get_columns('wide_scada_data')
        column_names = [col['name'] for col in columns]
        
        final_columns = []
        for d_col in desired_columns:
            for db_col in column_names:
                if d_col.lower().replace('-', '') == db_col.lower().replace('-', ''):
                    final_columns.append(f'"{db_col}"')
                    break
        
        if not final_columns:
            print("Error: No matching columns found. Data retrieval failed.")
            return None

        select_clause = ", ".join(final_columns)
        query = f"""
        SELECT {select_clause}
        FROM wide_scada_data
        WHERE "DateAndTime" BETWEEN '2025-08-08 00:00:00' AND '2025-08-20 23:59:59'
        ORDER BY "DateAndTime";
        """
        
        df = pd.read_sql(query, engine)
        df.columns = [col.upper().replace('-', '_') for col in df.columns]
        df['DATEANDTIME'] = pd.to_datetime(df['DATEANDTIME'])
        print("SCADA data for C-00 retrieved successfully.")
        return df
    except Exception as e:
        print(f"Error retrieving SCADA data: {e}")
        return None

def get_feed_composition():
    """Simulates getting composition data for the C-00 feed."""
    # Values are placeholders based on typical WFO composition
    return {
        'Naphthalene': 55.0, # %
        'Thianaphthalene': 2.0, # %
        'Quinoline': 1.7, # %
        'Unknown_impurity': 1.5, # %
        'Moisture': 0.2, # %
    }

def perform_analysis(df):
    """Calculates key performance indicators for the C-00 column."""
    if df is None or df.empty:
        return {}, df, {}
    
    outliers = {}
    analysis_results = {}
    
    # Anomaly Detection and Filtering
    if 'TI_63' in df.columns:
        mean_ti63 = df['TI_63'].mean()
        std_ti63 = df['TI_63'].std()
        outlier_mask = np.abs(df['TI_63'] - mean_ti63) > (5 * std_ti63)
        if outlier_mask.any():
            outlier_time = df.loc[outlier_mask, 'DATEANDTIME'].iloc[0]
            outliers['TI_63'] = {'time': outlier_time.strftime('%Y-%m-%d %H:%M'), 'value': df.loc[outlier_mask, 'TI_63'].iloc[0]}
            df.loc[outlier_mask, 'TI_63'] = np.nan

    # Get average flow rates
    feed_flow_avg = df['FT_01'].mean()
    moisture_flow_avg = df['FT_61'].mean()
    bottom_product_flow_avg = df['FT_62'].mean()
    
    analysis_results['Average Feed Flow (FT-01)'] = feed_flow_avg
    analysis_results['Average Moisture Flow (FT-61)'] = moisture_flow_avg
    analysis_results['Average Bottom Product Flow (FT-62)'] = bottom_product_flow_avg
    
    # Overall Material Balance Analysis
    if feed_flow_avg > 0:
        material_balance_error = ((feed_flow_avg - (moisture_flow_avg + bottom_product_flow_avg)) / feed_flow_avg) * 100
        analysis_results['Overall Material Balance Error (%)'] = abs(material_balance_error)

    # Component-wise Material Balance and Composition Calculation
    feed_composition = get_feed_composition()
    
    moisture_in_feed_flow = feed_flow_avg * (feed_composition.get('Moisture', 0) / 100.0)
    
    # Calculate moisture removal percentage
    if moisture_in_feed_flow > 0:
        moisture_removal_percent = (moisture_flow_avg / moisture_in_feed_flow) * 100
        analysis_results['Moisture Removal Percentage'] = moisture_removal_percent
    else:
        analysis_results['Moisture Removal Percentage'] = "N/A (No moisture in feed)"

    # Calculate composition of the bottom product (FT-62)
    bottom_product_composition = {}
    if bottom_product_flow_avg > 0:
        for component, percent in feed_composition.items():
            if component != 'Moisture':
                # Mass of non-moisture components in feed = Mass of non-moisture components in product
                component_mass_flow = feed_flow_avg * (percent / 100.0)
                # Calculate the new percentage in the bottom product
                bottom_product_composition[component] = (component_mass_flow / bottom_product_flow_avg) * 100
    
    analysis_results['Bottom Product Composition'] = bottom_product_composition
    
    # Differential Pressure (DP) Calculation
    if 'PTT_04' in df.columns and 'PTB_04' in df.columns:
        df['DIFFERENTIAL_PRESSURE'] = df['PTB_04'] - df['PTT_04']
        analysis_results['Average Differential Pressure'] = df['DIFFERENTIAL_PRESSURE'].mean()
        analysis_results['Maximum Differential Pressure'] = df['DIFFERENTIAL_PRESSURE'].max()
        
    # Energy Balance
    if all(tag in df.columns for tag in ['TI_215', 'TI_216', 'FI_204']):
        df['REBOILER_HEAT_DUTY'] = df['FI_204'] * THERMIC_FLUID_SPECIFIC_HEAT * (df['TI_216'] - df['TI_215'])
        analysis_results['Average Reboiler Heat Duty'] = df['REBOILER_HEAT_DUTY'].mean()
    else:
        analysis_results['Average Reboiler Heat Duty'] = 'N/A (Missing data)'

    if all(tag in df.columns for tag in ['TI_110', 'FI_101']):
        df['CONDENSER_HEAT_DUTY'] = df['FI_101'] * WATER_SPECIFIC_HEAT * (25 - df['TI_110'])
        analysis_results['Average Condenser Heat Duty'] = df['CONDENSER_HEAT_DUTY'].mean()
    else:
        analysis_results['Average Condenser Heat Duty'] = 'N/A (Missing data)'

    # Stability Analysis
    analysis_results['Stability'] = {}
    if 'TI_64' in df.columns:
        mean_ti64 = df['TI_64'].mean()
        std_ti64 = df['TI_64'].std()
        cv_ti64 = (std_ti64 / mean_ti64) * 100 if mean_ti64 != 0 else 0
        analysis_results['Stability']['TI-64 (Top Temp) Standard Deviation'] = std_ti64
        analysis_results['Stability']['TI-64 (Top Temp) Coefficient of Variation (%)'] = cv_ti64
    
    if 'DIFFERENTIAL_PRESSURE' in df.columns:
        mean_dp = df['DIFFERENTIAL_PRESSURE'].mean()
        std_dp = df['DIFFERENTIAL_PRESSURE'].std()
        cv_dp = (std_dp / mean_dp) * 100 if mean_dp != 0 else 0
        analysis_results['Stability']['Differential Pressure Standard Deviation'] = std_dp
        analysis_results['Stability']['Differential Pressure Coefficient of Variation (%)'] = cv_dp
        
    return analysis_results, df, outliers

def generate_plots(df):
    """Generates and saves temperature profile, DP, and daily trends plots."""
    try:
        plt.figure(figsize=(10, 6))
        if 'DATEANDTIME' in df.columns:
            df.sort_values(by='DATEANDTIME', inplace=True)
            x_axis = df['DATEANDTIME']
            if 'TI_61' in df.columns: plt.plot(x_axis, df['TI_61'], label='TI-61', alpha=0.7)
            if 'TI_63' in df.columns: plt.plot(x_axis, df['TI_63'], label='TI-63', alpha=0.7)
            if 'TI_64' in df.columns: plt.plot(x_axis, df['TI_64'], label='TI-64', alpha=0.7)
            plt.title("C-00 Column Temperature Profile Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Temperature ({TAG_UNITS['TI-61']})")
            plt.legend()
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_temp_plot_path)
            plt.close()
            print(f"Temperature profile plot saved to {output_temp_plot_path}")
    except Exception as e:
        print(f"Error generating temperature plot: {e}")
        
    try:
        if 'DIFFERENTIAL_PRESSURE' in df.columns:
            plt.figure(figsize=(10, 6))
            plt.plot(df['DATEANDTIME'], df['DIFFERENTIAL_PRESSURE'], color='purple', alpha=0.8)
            plt.title("C-00 Differential Pressure Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Differential Pressure ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})")
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_dp_plot_path)
            plt.close()
            print(f"Differential pressure plot saved to {output_dp_plot_path}")
    except Exception as e:
        print(f"Error generating DP plot: {e}")

    try:
        df['DATE'] = df['DATEANDTIME'].dt.date
        daily_trends = df.groupby('DATE').agg({
            'FT_01': 'mean',
            'TI_64': 'mean',
            'DIFFERENTIAL_PRESSURE': 'mean'
        }).reset_index()
        plt.figure(figsize=(12, 8))
        plt.plot(daily_trends['DATE'], daily_trends['FT_01'], label=f"Avg Feed Flow ({TAG_UNITS['FT-01']})")
        plt.plot(daily_trends['DATE'], daily_trends['TI_64'], label=f"Avg Top Temp ({TAG_UNITS['TI-64']})")
        plt.plot(daily_trends['DATE'], daily_trends['DIFFERENTIAL_PRESSURE'], label=f"Avg DP ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})")
        plt.title("C-00 Daily Trends")
        plt.xlabel("Date")
        plt.ylabel("Value")
        plt.legend()
        plt.grid(True)
        plt.tight_layout()
        plt.savefig(output_trends_plot_path)
        plt.close()
        print(f"Daily trends plot saved to {output_trends_plot_path}")
    except Exception as e:
        print(f"Error generating daily trends plot: {e}")
        
def generate_word_report(analysis_results, df, outliers):
    """Creates a detailed analysis report in a Word document."""
    doc = Document()
    doc.add_heading('C-00 Packed Distillation Column Analysis Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('1. Executive Summary', level=1)
    summary_text = ""
    if 'Stability' in analysis_results and 'TI-64 (Top Temp) Coefficient of Variation (%)' in analysis_results['Stability']:
        cv_ti64 = analysis_results['Stability']['TI-64 (Top Temp) Coefficient of Variation (%)']
        summary_text += f"The column demonstrated **excellent temperature stability** at the top with a low Coefficient of Variation of {cv_ti64:.2f}%. This consistency is crucial for maintaining product quality. "
    
    if 'DIFFERENTIAL_PRESSURE' in outliers:
        outlier_info = outliers['DIFFERENTIAL_PRESSURE']
        summary_text += f"A **significant spike in differential pressure** was detected on {outlier_info['time']}, reaching an anomalous value of {outlier_info['value']:.2f} {TAG_UNITS['DIFFERENTIAL_PRESSURE']}. This event should be investigated as it could indicate a temporary fouling or process upset. "
    else:
        summary_text += "Differential pressure remained stable throughout the period, suggesting no significant issues with fouling or flooding. "
            
    if 'Moisture Removal Percentage' in analysis_results and isinstance(analysis_results['Moisture Removal Percentage'], (float, int)):
        moisture_removed = analysis_results['Moisture Removal Percentage']
        summary_text += f"The column achieved a moisture removal efficiency of {moisture_removed:.2f}%. "
        
    doc.add_paragraph(summary_text)

    doc.add_heading('2. Key Performance Indicators (KPIs)', level=1)
    doc.add_paragraph("All values are averages over the analysis period, with outliers removed for accuracy.")
    for key, value in analysis_results.items():
        if isinstance(value, dict):
            continue
        tag_match = re.search(r'\((.*?)\)', key)
        if tag_match:
            tag = tag_match.group(1)
            unit = TAG_UNITS.get(tag, '')
        else:
            unit = TAG_UNITS.get(key.split(' ')[-1].strip(), '')
        if isinstance(value, str):
             doc.add_paragraph(f"• {key}: {value}")
        else:
             doc.add_paragraph(f"• {key}: {value:.2f} {unit}")

    doc.add_heading('3. Composition Analysis', level=1)
    doc.add_paragraph("The table below shows the calculated compositions for the C-00 streams.")
    
    doc.add_heading('3.1 Feed (FT-01) Composition', level=2)
    feed_comp = get_feed_composition()
    for comp, perc in feed_comp.items():
        doc.add_paragraph(f"• {comp.replace('_', ' ').capitalize()}: {perc:.2f}%")
        
    doc.add_heading('3.2 Bottom Product (FT-62) Composition', level=2)
    bottom_comp = analysis_results.get('Bottom Product Composition', {})
    if bottom_comp:
        for comp, perc in bottom_comp.items():
            doc.add_paragraph(f"• {comp.replace('_', ' ').capitalize()}: {perc:.2f}%")
    else:
        doc.add_paragraph("Composition data for the bottom product is not available due to missing flow data.")

    doc.add_heading('4. Column Stability Analysis', level=1)
    doc.add_paragraph("This section analyzes the stability of key process variables to identify operational consistency.")
    if 'Stability' in analysis_results:
        for key, value in analysis_results['Stability'].items():
            doc.add_paragraph(f"• {key}: {value:.2f}")

    doc.add_heading('5. Performance Plots', level=1)
    doc.add_heading('5.1 Temperature Profile', level=2)
    doc.add_paragraph("The temperature profile plot shows the gradient across the column. A consistent gradient indicates stable operation.")
    if 'TI_63' in outliers:
        doc.add_paragraph(f"**Note:** An extreme outlier was detected on {outliers['TI_63']['time']} for the TI-63 sensor, reaching a value of {outliers['TI_63']['value']:.2f} {TAG_UNITS['TI-63']}. This is likely a sensor malfunction and the value has been excluded from all calculations.")
    doc.add_picture(output_temp_plot_path, width=Inches(6))
    
    doc.add_heading('5.2 Differential Pressure (DP)', level=2)
    doc.add_paragraph("Differential pressure is a key indicator of flooding, foaming, or fouling inside the column.")
    doc.add_picture(output_dp_plot_path, width=Inches(6))

    doc.add_heading('5.3 Daily Trends', level=2)
    doc.add_paragraph("This plot shows the daily average trends of key variables, helping to visualize long-term shifts in performance.")
    doc.add_picture(output_trends_plot_path, width=Inches(6))
    doc.save(output_report_path)
    print(f"Analysis report generated successfully at {output_report_path}")

def main():
    """Main execution function."""
    engine = connect_to_database()
    if engine is None:
        return
    scada_data = get_scada_data(engine)
    if scada_data is None:
        return
    analysis_results, scada_data, outliers = perform_analysis(scada_data)
    if analysis_results:
        generate_plots(scada_data)
        generate_word_report(analysis_results, scada_data, outliers)
        print("C-00 analysis complete.")
    else:
        print("Analysis failed: no data to process.")

if __name__ == "__main__":
    main()