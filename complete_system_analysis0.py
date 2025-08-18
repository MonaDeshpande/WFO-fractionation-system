import psycopg2
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_report(df, filename, column_objectives):
    """Creates a Word document with analysis results and graphs."""
    doc = Document()
    doc.add_heading('Distillation Column Data Analysis Report', 0)
    
    # Check for thermic fluid temperature issue
    thermic_fluid_col = 'TI-72A'  # Assuming TI-72A is a representative tag for thermic fluid
    thermic_fluid_range = (325, 340)
    
    doc.add_heading('Overall System Analysis: Anthracene Purity', level=1)
    doc.add_paragraph('The overall aim of the system is to recover maximum pure anthracene from the bottom of Column C-03. This is a multi-stage process where each column plays a critical role in removing lighter components before the final purification step.')
    doc.add_paragraph('**Primary Factors Affecting Final Anthracene Purity:**')
    doc.add_paragraph('1. **Reflux Ratio:** For columns C-01 and C-02, a high reflux ratio is crucial to ensure efficient separation of naphthalene and light oil, preventing their carry-over to subsequent columns.')
    doc.add_paragraph('2. **Reboiler Temperature:** For all columns, maintaining stable and correct reboiler temperatures is essential. In C-03, specifically, the thermic fluid temperature must be in the correct range to vaporize residual light components.')
    
    doc.add_heading('Potential Operator Mistakes', level=2)
    doc.add_paragraph('Based on the provided data, a key area of concern is the reboiler thermic fluid temperature.')
    
    if thermic_fluid_col in df.columns:
        low_temp_points = df[df[thermic_fluid_col] < thermic_fluid_range[0]]
        high_temp_points = df[df[thermic_fluid_col] > thermic_fluid_range[1]]

        if not low_temp_points.empty:
            doc.add_paragraph(f"**🔴 Warning:** The {thermic_fluid_col} reading was below the specified range of {thermic_fluid_range[0]}°C. This may indicate insufficient heat input, which could compromise separation efficiency and lead to lower anthracene purity.")
            doc.add_paragraph(f"Lowest recorded temperature: {low_temp_points[thermic_fluid_col].min():.2f}°C on {low_temp_points['datetime'].iloc[0]}")
        
        if not high_temp_points.empty:
            doc.add_paragraph(f"**🔴 Warning:** The {thermic_fluid_col} reading was above the specified range of {thermic_fluid_range[1]}°C. This could cause thermal degradation of the product or unnecessary energy consumption.")
            doc.add_paragraph(f"Highest recorded temperature: {high_temp_points[thermic_fluid_col].max():.2f}°C on {high_temp_points['datetime'].iloc[0]}")
            
    doc.add_page_break()

    # Analyze each column
    for column_name in ['C-00', 'C-01', 'C-02', 'C-03']:
        doc.add_heading(f'Analysis for {column_name}', level=1)
        doc.add_paragraph(column_objectives[column_name]['purpose'])

        if column_name == 'C-00':
            if 'FT-01' in df.columns and 'FI-61' in df.columns:
                doc.add_paragraph('**Key Operating Ranges for C-00:**')
                doc.add_paragraph(f'Initial Feed Rate (FT-01): Average {df["FT-01"].mean():.2f}, Range [{df["FT-01"].min():.2f}, {df["FT-01"].max():.2f}]')
                doc.add_paragraph(f'Water Removal Rate (FI-61): Average {df["FI-61"].mean():.2f}, Range [{df["FI-61"].min():.2f}, {df["FI-61"].max():.2f}]')
        
        else:
            reflux_flow, top_product_flow, feed_temp_col, temp_cols = '', '', '', []
            if column_name == 'C-01':
                reflux_flow = 'FT-08'
                top_product_flow = 'FT-02'
                feed_temp_col = 'TI-02'
                temp_cols = ['TI-03', 'TI-04', 'TI-05', 'TI-06']
            elif column_name == 'C-02':
                reflux_flow = 'FT-09'
                top_product_flow = 'FT-03'
                feed_temp_col = 'TI-11'
                temp_cols = ['TI-13', 'TI-14', 'TI-15', 'TI-16', 'TI-17', 'TI-18', 'TI-19', 'TI-20', 'TI-21', 'TI-22', 'TI-23', 'TI-24', 'TI-25']
            elif column_name == 'C-03':
                reflux_flow = 'FT-10'
                top_product_flow = 'FT-04'
                feed_temp_col = 'TI-30'
                temp_cols = ['TI-31', 'TI-32', 'TI-33', 'TI-34', 'TI-35', 'TI-36', 'TI-37', 'TI-38', 'TI-39', 'TI-40']

            if reflux_flow in df.columns and top_product_flow in df.columns:
                df['reflux_ratio'] = df[reflux_flow] / df[top_product_flow]
                
                doc.add_paragraph(f"**Key Operating Ranges for {column_name}:**")
                doc.add_paragraph(f"Average Reflux Ratio: {df['reflux_ratio'].mean():.2f}")
                doc.add_paragraph(f"Average Feed Temp ({feed_temp_col}): {df[feed_temp_col].mean():.2f}°C")

                doc.add_heading('Process Variable Trends', level=2)
                fig, axes = plt.subplots(nrows=3, ncols=1, figsize=(10, 15))
                
                axes[0].plot(df['datetime'], df[feed_temp_col])
                axes[0].set_title(f'Feed Temperature ({feed_temp_col}) vs. Time')
                axes[0].set_xlabel('Time')
                axes[0].set_ylabel('Temperature (°C)')
                
                axes[1].plot(df['datetime'], df['reflux_ratio'])
                axes[1].set_title('Reflux Ratio vs. Time')
                axes[1].set_xlabel('Time')
                axes[1].set_ylabel('Reflux Ratio')
                
                for col in temp_cols:
                    axes[2].plot(df['datetime'], df[col], label=col)
                axes[2].set_title('Column Temperatures vs. Time')
                axes[2].set_xlabel('Time')
                axes[2].set_ylabel('Temperature (°C)')
                axes[2].legend()
                
                plt.tight_layout()
                plt_filename = f'plot_{column_name}_process.png'
                plt.savefig(plt_filename)
                doc.add_picture(plt_filename, width=Inches(6))

                doc.add_heading('Factors Affecting Purity', level=2)
                if column_name == 'C-01':
                    doc.add_paragraph('**Goal: Anthracene oil with <2% naphthalene.** The most critical factor is **reboiler temperature and flow rate**, as sufficient heat is needed to strip naphthalene from the bottom product.')
                elif column_name == 'C-02':
                    doc.add_paragraph('**Goal: Light oil with <15% naphthalene.** The most critical factor is **reflux ratio**. A higher reflux ratio improves separation, ensuring naphthalene is effectively concentrated in the top product.')
                elif column_name == 'C-03':
                    doc.add_paragraph('**Goal: Maximum naphthalene in top product and pure anthracene in bottom.** The most critical factors are **reboiler temperature (325-340°C)** and the **column bottom temperature**. Maintaining these temperatures is essential to vaporize any remaining light components, resulting in a purer anthracene product at the bottom.')
        
        doc.add_page_break()

    doc.save(filename)
    print(f"Report saved as {filename}")

def get_data_from_db(start_date, end_date):
    """
    Connects to the PostgreSQL database and fetches the data.
    IMPORTANT: You must fill in your database credentials and table names.
    """
    conn = None
    df = pd.DataFrame()
    try:
        conn = psycopg2.connect(
            dbname="your_db_name",
            user="your_user_name",
            password="your_password",
            host="your_host",
            port="your_port"
        )
        print("Database connection successful.")
        
        query = f"""
        SELECT 
            t1."DateAndTime" AS datetime,
            t1."PTT-04", t1."TI-61", t1."TI-63", t1."TI-64", t1."TI-110", t1."TI-112", t1."TI-01", t1."FT-02",
            t2."TI-30", t2."TI-207", t2."TI-208", t2."TI-72A", t2."TI-72B", t2."TI-106", t2."TI-209", t2."TI-210",
            t3."TI-04", t3."TI-05", t3."TI-06", t3."TI-07", t3."TI-08", t3."TI-09", t3."TI-10", t3."TI-11",
            t4."FT-01", t4."FT-62", t4."FT-61", t4."LT-61", t4."PTB-04", t4."PTT-04",
            t5."TI-31", t5."TI-32", t5."TI-33", t5."TI-34", t5."TI-35", t5."TI-36", t5."TI-37", t5."TI-38",
            t6."TI-22", t6."TI-23", t6."TI-24", t6."TI-25", t6."TI-26", t6."TI-27", t6."TI-28", t6."TI-29",
            t7."TI-210" AS "TI-210_copy", t7."FT-10", t7."FT-07", t7."FT-04", t7."LT-05", t7."LT-06", t7."PTB-03", t7."PTT-03",
            t8."TI-39", t8."TI-40", t8."TI-41", t8."TI-42", t8."TI-45", t8."TI-54", t8."TI-55", t8."TI-42A",
            t9."FT-09", t9."FT-06", t9."LT-03", t9."LT-04", t9."PTB-02", t9."PTT-02", t9."TI-13",
            t10."TI-14", t10."TI-15", t10."TI-16", t10."TI-17", t10."TI-18", t10."TI-19", t10."TI-20", t10."TI-21"
        FROM 
            "table_for_IMG-20250814-WA0115.jpg" AS t1
            JOIN "table_for_IMG-20250814-WA0118.jpg" AS t2 ON t1."row_id" = t2."row_id"
            JOIN "table_for_IMG-20250814-WA0116.jpg" AS t3 ON t1."row_id" = t3."row_id"
            JOIN "table_for_IMG-20250814-WA0114.jpg" AS t4 ON t1."row_id" = t4."row_id"
            JOIN "table_for_IMG-20250814-WA0120.jpg" AS t5 ON t1."row_id" = t5."row_id"
            JOIN "table_for_IMG-20250814-WA0117.jpg" AS t6 ON t1."row_id" = t6."row_id"
            JOIN "table_for_IMG-20250814-WA0119.jpg" AS t7 ON t1."row_id" = t7."row_id"
            JOIN "table_for_IMG-20250814-WA0121.jpg" AS t8 ON t1."row_id" = t8."row_id"
            JOIN "table_for_IMG-20250814-WA0124.jpg" AS t9 ON t1."row_id" = t9."row_id"
            JOIN "table_for_IMG-20250814-WA0123.jpg" AS t10 ON t1."row_id" = t10."row_id"
        WHERE 
            "DateAndTime" BETWEEN '{start_date}' AND '{end_date}'
        ORDER BY 
            "DateAndTime";
        """
        df = pd.read_sql(query, conn)
        df['datetime'] = pd.to_datetime(df['datetime'])

    except (Exception, psycopg2.DatabaseError) as error:
        print(error)
    finally:
        if conn is not None:
            conn.close()
    return df

if __name__ == '__main__':
    # --- USER INPUT SECTION ---
    # Database connection details
    db_credentials = {
        'dbname': 'your_db_name',
        'user': 'your_user_name',
        'password': 'your_password',
        'host': 'your_host',
        'port': 'your_port'
    }

    # Time range for analysis
    start_time = '2025-08-08 00:00:00'
    end_time = '2025-08-18 23:59:59'

    # Output file name
    output_filename = 'Distillation_Column_Report.docx'

    # Column objectives and key parameters
    column_objectives = {
        'C-00': {
            'purpose': 'This column aims to perform dehydration of the feed.',
        },
        'C-01': {
            'purpose': 'To achieve a bottom product (Anthracene Oil) with less than 2% naphthalene oil.',
        },
        'C-02': {
            'purpose': 'To achieve a top product (Light Oil) with less than 15% naphthalene oil.',
        },
        'C-03': {
            'purpose': 'To achieve a top product (Naphthalene Oil) with maximum naphthalene percentage, and a bottom product with pure anthracene.',
        }
    }
    # --------------------------

    # Step 1: Get data from PostgreSQL
    full_df = get_data_from_db(start_time, end_time)
    
    if not full_df.empty:
        # Step 2: Create the Word report with analysis
        create_word_report(full_df, output_filename, column_objectives)
    else:
        print("No data found in the specified time range. Please check your date range and database connection.")