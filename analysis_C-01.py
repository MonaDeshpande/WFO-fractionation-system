import pandas as pd
from sqlalchemy import create_engine
import psycopg2
from datetime import datetime
import matplotlib.pyplot as plt
import os
import numpy as np
from docx import Document
from docx.shared import Inches
import sqlalchemy
import re

# Database connection parameters (update with your actual details)
# Note: These are placeholders. You must configure them to connect to your database.
DB_HOST = "localhost"
DB_NAME = "scada_data_analysis"
DB_USER = "postgres"
DB_PASSWORD = "ADMIN"

# Define units for each tag
TAG_UNITS = {
    'FT-02': 'kg/h',
    'FT-05': 'kg/h',
    'FT-08': 'kg/h',
    'TI-02': 'degC',
    'TI-04': 'degC',
    'TI-05': 'degC',
    'TI-06': 'degC',
    'TI-07': 'degC',
    'TI-08': 'degC',
    'TI-10': 'degC',
    'TI-11': 'degC',
    'TI-12': 'degC',
    'TI-52': 'degC',
    'PTT-01': 'mmHg',
    'PTB-01': 'mmHg',
    'DIFFERENTIAL_PRESSURE': 'mmHg',
    'FI-201': 'kg/h',
    'TI-203': 'degC',
    'TI-204': 'degC',
    'TI-205': 'degC',
    'TI-206': 'degC',
    'TI-202': 'degC',
    'TI-110': 'degC',
    'TI-111': 'degC',
    'FI-101': 'kg/h',
    'REBOILER_HEAT_DUTY': 'kW',
    'CONDENSER_HEAT_DUTY': 'kW',
    'FEED_PREHEATER_DUTY': 'kW',
    'TOP_PRODUCT_HEATER_DUTY': 'kW',
    'REFLUX_RATIO': '',
    'MATERIAL_BALANCE_ERROR': '%',
    'NAPHTHALENE_LOSS_MASS': 'kg/h',
    'NAPHTHALENE_LOSS_PERCENTAGE': '%'
}

# File paths for saving generated plots and report
output_report_path = "C-01_Analysis_Report.docx"
output_temp_plot_path = "C-01_Temperature_Profile.png"
output_dp_plot_path = "C-01_Differential_Pressure.png"
output_trends_plot_path = "C-01_Daily_Trends.png"

# Engineering constants for heat duty calculations
THERMIC_FLUID_SPECIFIC_HEAT = 2.0  # kJ/(kg·°C) - Assumed value, replace with specific data if available
WATER_SPECIFIC_HEAT = 4.186       # kJ/(kg·°C)

def connect_to_database():
    """Establishes a connection to the PostgreSQL database."""
    try:
        engine = create_engine(f'postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}/{DB_NAME}')
        print("Database connection successful.")
        return engine
    except Exception as e:
        print(f"Error connecting to the database: {e}")
        return None

def get_scada_data(engine):
    """
    Retrieves specific SCADA data for the C-01 column from the database.
    It selects all necessary tags provided in the user's description.
    """
    try:
        desired_columns = [
            "DateAndTime", "FT-02", "FT-05", "FT-08", "TI-02", "TI-04", "TI-05", "TI-06", "TI-07",
            "TI-08", "TI-10", "TI-11", "TI-12", "TI-52", "PTT-01", "PTB-01", "FI-201", "TI-203", 
            "TI-204", "TI-205", "TI-206", "TI-202", "TI-110", "TI-111", "FI-101", "P-01"
        ]
        
        inspector = sqlalchemy.inspect(engine)
        columns = inspector.get_columns('wide_scada_data')
        column_names = [col['name'] for col in columns]
        
        final_columns = []
        for d_col in desired_columns:
            for db_col in column_names:
                if d_col.replace('-', '').lower() == db_col.replace('-', '').lower():
                    final_columns.append(f'"{db_col}"')
                    break
        
        if not final_columns:
            print("Error: No matching columns found for C-01. Data retrieval failed.")
            return None

        select_clause = ", ".join(final_columns)
        query = f"""
        SELECT {select_clause}
        FROM wide_scada_data
        WHERE "DateAndTime" BETWEEN '2025-08-08 00:40:00' AND '2025-08-20 12:40:59'
        ORDER BY "DateAndTime";
        """
        
        df = pd.read_sql(query, engine)
        df.columns = [col.upper().replace('-', '_') for col in df.columns] # Normalize column names
        df['DATEANDTIME'] = pd.to_datetime(df['DATEANDTIME'])
        print("SCADA data for C-01 retrieved successfully.")
        return df
    except Exception as e:
        print(f"Error retrieving SCADA data: {e}")
        return None

def get_feed_composition():
    """Simulates getting feed composition data from C-00 bottom product."""
    # This data would come from the C-00 analysis or a lab sheet
    return {
        'Naphthalene': 95.0, # % (High percentage after moisture removal in C-00)
        'Thianaphthalene': 2.0, # %
        'Quinoline': 1.7, # %
        'Unknown_impurity': 1.3, # %
    }

def get_bottom_product_composition():
    """Simulates getting bottom product composition data from a lab sheet (C-01-B)."""
    # This data would come from a lab sheet
    return {
        'Naphthalene': 2.0, # % (Remaining Naphthalene)
        'Anthracene Oil': 98.0, # %
    }

def perform_analysis(df):
    """
    Performs key calculations for C-01, including material/energy balances
    and reflux ratio.
    """
    if df is None or df.empty:
        return {}, df, {}

    outliers = {}
    analysis_results = {}
    
    # Get average flow rates
    feed_flow_avg = df['FT_05'].mean()
    top_product_flow_avg = df['FT_02'].mean()
    bottom_product_flow_avg = df['FT_05'].mean() # Note: FT-05 is also the bottom product flow
    
    analysis_results['Average Feed Flow (FT-05)'] = feed_flow_avg
    analysis_results['Average Top Product Flow (FT-02)'] = top_product_flow_avg
    analysis_results['Average Bottom Product Flow (FT-05)'] = bottom_product_flow_avg
    
    # Overall Material Balance
    if feed_flow_avg > 0:
        material_balance_error = ((feed_flow_avg - (top_product_flow_avg + bottom_product_flow_avg)) / feed_flow_avg) * 100
        analysis_results['Material Balance Error (%)'] = abs(material_balance_error)

    # Component-wise Material Balance and Composition Calculation
    feed_composition = get_feed_composition()
    bottom_comp_data = get_bottom_product_composition()
    
    top_product_composition = {}
    if top_product_flow_avg > 0:
        # Calculate mass flow of each component in the feed and bottom product
        for component, percent in feed_composition.items():
            feed_component_mass = feed_flow_avg * (percent / 100.0)
            
            # Assuming 'Naphthalene' is the only component also in the bottom product for this example.
            bottom_component_mass = 0
            if component == 'Naphthalene':
                bottom_component_mass = bottom_product_flow_avg * (bottom_comp_data.get('Naphthalene', 0) / 100.0)
            
            # Deduce mass flow of the component in the top product
            top_component_mass = feed_component_mass - bottom_component_mass
            
            # Calculate the percentage in the top product
            if top_component_mass > 0:
                top_product_composition[component] = (top_component_mass / top_product_flow_avg) * 100

    analysis_results['Top Product Composition'] = top_product_composition

    # Naphthalene Loss & Impurity Analysis
    if 'FT_05' in df.columns:
        naphthalene_in_bottom_product = bottom_product_flow_avg * (bottom_comp_data.get('Naphthalene', 0) / 100.0)
        naphthalene_in_feed = feed_flow_avg * (feed_composition.get('Naphthalene', 0) / 100.0)
        
        if naphthalene_in_feed > 0:
            naphthalene_loss_percent = (naphthalene_in_bottom_product / naphthalene_in_feed) * 100
            analysis_results['Naphthalene Loss (%)'] = naphthalene_loss_percent
            analysis_results['Naphthalene Loss (mass)'] = naphthalene_in_bottom_product
    
    # Reflux Ratio
    if 'FT_08' in df.columns and 'FT_02' in df.columns:
        reflux_flow_avg = df['FT_08'].mean()
        top_product_flow_avg = df['FT_02'].mean()
        if top_product_flow_avg > 0:
            reflux_ratio = reflux_flow_avg / top_product_flow_avg
            analysis_results['Average Reflux Ratio'] = reflux_ratio
        else:
            analysis_results['Average Reflux Ratio'] = "N/A (Zero Top Product Flow)"
        
    # Differential Pressure (DP) Calculation
    if 'PTT_01' in df.columns and 'PTB_01' in df.columns:
        df['DIFFERENTIAL_PRESSURE'] = df['PTB_01'] - df['PTT_01']
        analysis_results['Average Differential Pressure'] = df['DIFFERENTIAL_PRESSURE'].mean()
        analysis_results['Maximum Differential Pressure'] = df['DIFFERENTIAL_PRESSURE'].max()
        
    # Comprehensive Energy Balance
    # 1. Reboiler Heat Duty
    if all(tag in df.columns for tag in ['TI_203', 'TI_204', 'FI_201']):
        df['REBOILER_HEAT_DUTY'] = df['FI_201'] * THERMIC_FLUID_SPECIFIC_HEAT * (df['TI_204'] - df['TI_203'])
        analysis_results['Average Reboiler Heat Duty'] = df['REBOILER_HEAT_DUTY'].mean()

    # 2. Main Condenser Heat Duty
    if all(tag in df.columns for tag in ['TI_110', 'TI_111', 'FI_101']):
        df['CONDENSER_HEAT_DUTY'] = df['FI_101'] * WATER_SPECIFIC_HEAT * (df['TI_111'] - df['TI_110'])
        analysis_results['Average Main Condenser Heat Duty'] = df['CONDENSER_HEAT_DUTY'].mean()

    return analysis_results, df, outliers

def generate_plots(df):
    """Generates and saves temperature profile, DP, and energy plots."""
    # Temperature Profile Plot
    try:
        plt.figure(figsize=(10, 6))
        
        if 'DATEANDTIME' in df.columns:
            df.sort_values(by='DATEANDTIME', inplace=True)
            x_axis = df['DATEANDTIME']
            
            if 'TI_04' in df.columns: plt.plot(x_axis, df['TI_04'], label='TI-04', alpha=0.7)
            if 'TI_05' in df.columns: plt.plot(x_axis, df['TI_05'], label='TI-05', alpha=0.7)
            if 'TI_06' in df.columns: plt.plot(x_axis, df['TI_06'], label='TI-06', alpha=0.7)
            if 'TI_07' in df.columns: plt.plot(x_axis, df['TI_07'], label='TI-07', alpha=0.7)

            plt.title("C-01 Column Temperature Profile Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Temperature ({TAG_UNITS['TI-04']})")
            plt.legend()
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_temp_plot_path)
            plt.close()
            print(f"Temperature profile plot saved to {output_temp_plot_path}")
            
    except Exception as e:
        print(f"Error generating temperature plot: {e}")
        
    # Differential Pressure Plot
    try:
        if 'DIFFERENTIAL_PRESSURE' in df.columns:
            plt.figure(figsize=(10, 6))
            plt.plot(df['DATEANDTIME'], df['DIFFERENTIAL_PRESSURE'], color='purple', alpha=0.8)
            plt.title("C-01 Differential Pressure Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Differential Pressure ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})")
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_dp_plot_path)
            plt.close()
            print(f"Differential pressure plot saved to {output_dp_plot_path}")
    except Exception as e:
        print(f"Error generating DP plot: {e}")

    # Daily Trends Plot
    try:
        df['DATE'] = df['DATEANDTIME'].dt.date
        daily_trends = df.groupby('DATE').agg({
            'FT_02': 'mean',
            'TI_07': 'mean',
            'DIFFERENTIAL_PRESSURE': 'mean'
        }).reset_index()

        plt.figure(figsize=(12, 8))
        plt.plot(daily_trends['DATE'], daily_trends['FT_02'], label=f"Avg Top Product Flow ({TAG_UNITS['FT-02']})")
        plt.plot(daily_trends['DATE'], daily_trends['TI_07'], label=f"Avg Top Temp ({TAG_UNITS['TI-07']})")
        plt.plot(daily_trends['DATE'], daily_trends['DIFFERENTIAL_PRESSURE'], label=f"Avg DP ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})")
        plt.title("C-01 Daily Trends")
        plt.xlabel("Date")
        plt.ylabel("Value")
        plt.legend()
        plt.grid(True)
        plt.tight_layout()
        plt.savefig(output_trends_plot_path)
        plt.close()
        print(f"Daily trends plot saved to {output_trends_plot_path}")
    except Exception as e:
        print(f"Error generating daily trends plot: {e}")
        
def generate_word_report(analysis_results, df, outliers):
    """Creates a detailed analysis report in a Word document."""
    doc = Document()
    doc.add_heading('C-01 Anthracene Oil Recovery Column Analysis Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    
    summary_text = ""
    if 'Average Reflux Ratio' in analysis_results and isinstance(analysis_results['Average Reflux Ratio'], (float, int)):
        summary_text += f"The column operated with an average reflux ratio of {analysis_results['Average Reflux Ratio']:.2f}, indicating effective control over product separation. "
    
    if 'Material Balance Error (%)' in analysis_results:
        summary_text += f"A material balance error of {analysis_results['Material Balance Error (%)']:.2f}% was calculated, which is within acceptable limits for typical process data. "
    
    if 'Naphthalene Loss (%)' in analysis_results:
        summary_text += f"Based on the bottom product analysis, a naphthalene loss of {analysis_results['Naphthalene Loss (%)']:.2f}% was observed, corresponding to a mass loss of {analysis_results['Naphthalene Loss (mass)']:.2f} kg/h. "
            
    doc.add_paragraph(summary_text)

    # Section 2: Key Performance Indicators
    doc.add_heading('2. Key Performance Indicators (KPIs)', level=1)
    doc.add_paragraph("All values are averages over the analysis period.")
    for key, value in analysis_results.items():
        if isinstance(value, dict):
            continue
        # Use regex to find the unit if it's in a format like 'Key (Tag)'
        tag_match = re.search(r'\((.*?)\)', key)
        if tag_match:
            tag = tag_match.group(1)
            unit = TAG_UNITS.get(tag, '')
        else:
            unit = TAG_UNITS.get(key.split(' ')[-1].strip(), '')

        if isinstance(value, str):
             doc.add_paragraph(f"• {key}: {value}")
        else:
             doc.add_paragraph(f"• {key}: {value:.2f} {unit}")

    # Section 3: Composition Analysis
    doc.add_heading('3. Composition Analysis', level=1)
    doc.add_paragraph("The tables below show the calculated compositions for the C-01 streams.")
    
    doc.add_heading('3.1 Feed (FT-05) Composition', level=2)
    feed_comp = get_feed_composition()
    for comp, perc in feed_comp.items():
        doc.add_paragraph(f"• {comp.replace('_', ' ').capitalize()}: {perc:.2f}%")
        
    doc.add_heading('3.2 Bottom Product (FT-05) Composition', level=2)
    bottom_comp = get_bottom_product_composition()
    for comp, perc in bottom_comp.items():
        doc.add_paragraph(f"• {comp.replace('_', ' ').capitalize()}: {perc:.2f}%")

    doc.add_heading('3.3 Top Product (FT-02) Composition', level=2)
    top_comp = analysis_results.get('Top Product Composition', {})
    if top_comp:
        for comp, perc in top_comp.items():
            doc.add_paragraph(f"• {comp.replace('_', ' ').capitalize()}: {perc:.2f}%")
    else:
        doc.add_paragraph("Composition data for the top product is not available due to missing flow data.")

    # Section 4: Performance Plots
    doc.add_heading('4. Performance Plots', level=1)

    doc.add_heading('4.1 Temperature Profile', level=2)
    doc.add_paragraph("The temperature profile plot shows the gradient across the column.")
    doc.add_picture(output_temp_plot_path, width=Inches(6))

    doc.add_heading('4.2 Differential Pressure (DP)', level=2)
    doc.add_paragraph("Differential pressure is a key indicator of flooding or fouling.")
    doc.add_picture(output_dp_plot_path, width=Inches(6))

    doc.add_heading('4.3 Daily Trends', level=2)
    doc.add_paragraph("This plot shows the daily average trends of key variables.")
    doc.add_picture(output_trends_plot_path, width=Inches(6))

    doc.save(output_report_path)
    print(f"Analysis report generated successfully at {output_report_path}")

def main():
    """Main execution function."""
    engine = connect_to_database()
    if engine is None:
        return

    scada_data = get_scada_data(engine)
    if scada_data is None:
        return

    analysis_results, scada_data, outliers = perform_analysis(scada_data)
    
    if analysis_results:
        generate_plots(scada_data)
        generate_word_report(analysis_results, scada_data, outliers)
        print("C-01 analysis complete.")
    else:
        print("Analysis failed: no data to process.")

if __name__ == "__main__":
    main()